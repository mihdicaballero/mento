import pytest
import pandas as pd
from docx import Document
from typing import List, Any
from docx.shared import Cm
import matplotlib
from mento.results import Formatter, TablePrinter, DocumentBuilder, configure_plot_settings

matplotlib.use("Agg")  # Use non-GUI backend for testing
import matplotlib.pyplot as plt

# Disable LaTeX for testing to avoid CI/CD issues
plt.rcParams["text.usetex"] = False


@pytest.fixture
def formatter() -> Formatter:
    return Formatter()


@pytest.fixture
def table_printer() -> TablePrinter:
    return TablePrinter()


@pytest.fixture
def table_printer_with_title() -> TablePrinter:
    return TablePrinter(title="Test Table")


@pytest.fixture
def document_builder() -> DocumentBuilder:
    return DocumentBuilder(title="Test Document")


# --- Tests for configure_plot_settings function ---


def test_configure_plot_settings() -> None:
    """Test that configure_plot_settings runs without error."""
    # Temporarily disable usetex for testing
    original_usetex = plt.rcParams.get("text.usetex", False)
    plt.rcParams["text.usetex"] = False

    configure_plot_settings()

    # Verify some settings were applied (but not usetex in CI)
    assert plt.rcParams["axes.titlesize"] == 12
    assert plt.rcParams["axes.labelsize"] == 12
    assert plt.rcParams["legend.fontsize"] == 12

    # Restore original setting
    plt.rcParams["text.usetex"] = original_usetex


# --- Tests for Formatter class ---


def test_formatter_initialization(formatter: Formatter) -> None:
    """Test Formatter initializes with correct attributes."""
    assert formatter.green == "#439b00"
    assert formatter.red == "#d43e36"
    assert formatter.yellow == "#efc200"
    assert formatter.mid_value == 0.95
    assert formatter.max_value == 1


def test_formatter_DCR(formatter: Formatter) -> None:
    """Test DCR formatting with different values."""
    assert formatter.DCR(0.85) == "$\\color{#439b00}{\\text{DCR}=0.85}$"
    assert formatter.DCR(0.95) == "$\\color{#efc200}{\\text{DCR}=0.95}$"
    assert formatter.DCR(1.0) == "$\\color{#efc200}{\\text{DCR}=1.0}$"
    assert formatter.DCR(1.2) == "$\\color{#d43e36}{\\text{DCR}=1.2}$"


def test_formatter_DCR_value(formatter: Formatter) -> None:
    """Test DCR_value formatting with different values."""
    assert formatter.DCR_value(0.85) == "$\\color{#439b00}{0.85}$"
    assert formatter.DCR_value(0.95) == "$\\color{#efc200}{0.95}$"
    assert formatter.DCR_value(1.0) == "$\\color{#efc200}{1.0}$"
    assert formatter.DCR_value(1.2) == "$\\color{#d43e36}{1.2}$"


def test_formatter_is_lower(formatter: Formatter) -> None:
    """Test is_lower comparison."""
    assert formatter.is_lower(0.8, 1.0) == "$\\color{#439b00}{\\, \\checkmark}$"
    assert formatter.is_lower(1.2, 1.0) == "$\\color{#d43e36}{\\, \\times}$"
    assert formatter.is_lower(1.0, 1.0) == "$\\color{#d43e36}{\\, \\times}$"


def test_formatter_is_greater(formatter: Formatter) -> None:
    """Test is_greater comparison."""
    assert formatter.is_greater(1.2, 1.0) == "$\\color{#439b00}{\\, \\checkmark}$"
    assert formatter.is_greater(0.8, 1.0) == "$\\color{#d43e36}{\\, \\times}$"
    assert formatter.is_greater(1.0, 1.0) == "$\\color{#d43e36}{\\, \\times}$"


def test_formatter_DCR_value_df_with_valid_values(formatter: Formatter) -> None:
    """Test DCR_value_df with valid numeric values."""
    assert formatter.DCR_value_df(0.85) == "color: #439b00"
    assert formatter.DCR_value_df(0.95) == "color: #efc200"
    assert formatter.DCR_value_df(1.0) == "color: #efc200"
    assert formatter.DCR_value_df(1.2) == "color: #d43e36"


def test_formatter_DCR_value_df_with_none(formatter: Formatter) -> None:
    """Test DCR_value_df returns empty string for None."""
    assert formatter.DCR_value_df(None) == ""


def test_formatter_DCR_value_df_with_invalid_type(formatter: Formatter) -> None:
    """Test DCR_value_df returns empty string for invalid types."""
    assert formatter.DCR_value_df("invalid") == ""  # type: ignore
    assert formatter.DCR_value_df([1, 2, 3]) == ""  # type: ignore


def test_formatter_apply_DCR_style_with_float(formatter: Formatter) -> None:
    """Test apply_DCR_style with float value."""
    assert formatter.apply_DCR_style(0.856789) == "color: #439b00"
    assert formatter.apply_DCR_style(1.234) == "color: #d43e36"


def test_formatter_apply_DCR_style_with_int(formatter: Formatter) -> None:
    """Test apply_DCR_style with integer value."""
    assert formatter.apply_DCR_style(0) == "color: #439b00"
    assert formatter.apply_DCR_style(2) == "color: #d43e36"


def test_formatter_apply_DCR_style_with_non_numeric(formatter: Formatter) -> None:
    """Test apply_DCR_style with non-numeric value."""
    result = formatter.apply_DCR_style("text")  # type: ignore
    assert result == ""


def test_formatter_color_DCR_df(formatter: Formatter) -> None:
    """Test color_DCR_df styling."""
    df = pd.DataFrame({"DCR": [0.85, 0.95, 1.2], "Other": [1, 2, 3]})
    styled_df = formatter.color_DCR_df(df, ["DCR"])

    # Check that a Styler object is returned
    assert styled_df is not None
    assert hasattr(styled_df, "data")  # Styler objects have a 'data' attribute


# --- Tests for TablePrinter class ---


def test_table_printer_initialization_no_title() -> None:
    """Test TablePrinter initializes without title."""
    printer = TablePrinter()
    assert printer.title is None


def test_table_printer_initialization_with_title(table_printer_with_title: TablePrinter) -> None:
    """Test TablePrinter initializes with title."""
    assert table_printer_with_title.title == "Test Table"


def test_table_printer_print_table_data(table_printer: TablePrinter, capsys: Any) -> None:
    """Test print_table_data outputs correctly."""
    data = {
        "Column1": ["Value1", "Value2"],
        "Column2": ["a", "b"],
        "Column3": [5.5, 6.7],
        "Column4": ["cm", "mm"],
    }
    headers = ["Column1", "Column2", "Column3", "Column4"]
    result = table_printer.print_table_data(data, headers)

    assert result is None  # Method returns None

    captured = capsys.readouterr()
    assert "Value1" in captured.out
    assert "Value2" in captured.out
    assert "5.5" in captured.out
    assert "6.7" in captured.out


def test_table_printer_print_table_data_with_custom_format(table_printer: TablePrinter, capsys: Any) -> None:
    """Test print_table_data with custom tablefmt."""
    # Need 4 columns to match the colalign tuple in the method
    data = {"Col1": ["A", "B"], "Col2": [1, 2], "Col3": [3, 4], "Col4": ["x", "y"]}
    headers = ["Col1", "Col2", "Col3", "Col4"]
    table_printer.print_table_data(data, headers, tablefmt="grid", numalign="center")

    captured = capsys.readouterr()
    assert "A" in captured.out
    assert "B" in captured.out


def test_table_printer_print_table_min_max(table_printer_with_title: TablePrinter, capsys: Any) -> None:
    """Test print_table_min_max outputs correctly."""
    data = {
        "Item": ["Test1", "Test2"],
        "Min": [1, 2],
        "Max": [10, 20],
        "Avg": [5, 10],
        "Val1": [3, 6],
        "Val2": [7, 14],
        "Unit": ["m", "m"],
    }
    headers = ["Item", "Min", "Max", "Avg", "Val1", "Val2", "Unit"]
    table_printer_with_title.print_table_min_max(data, headers)

    captured = capsys.readouterr()
    assert "Test Table" in captured.out
    assert "Test1" in captured.out
    assert "Test2" in captured.out


def test_table_printer_print_table_min_max_no_title(table_printer: TablePrinter, capsys: Any) -> None:
    """Test print_table_min_max without title."""
    data = {"Item": ["Test"], "Min": [1], "Max": [2], "A": [3], "B": [4], "C": [5], "Unit": ["m"]}
    headers = ["Item", "Min", "Max", "A", "B", "C", "Unit"]
    table_printer.print_table_min_max(data, headers)

    captured = capsys.readouterr()
    # Should not have a title line
    assert "=======" not in captured.out or "Test Table" not in captured.out


# --- Tests for DocumentBuilder class ---


def test_document_builder_initialization(document_builder: DocumentBuilder) -> None:
    """Test DocumentBuilder initializes correctly."""
    assert document_builder.title == "Test Document"
    assert document_builder.font_name == "Lato"
    assert document_builder.font_size == 9
    assert isinstance(document_builder.doc, type(Document()))


def test_document_builder_custom_font() -> None:
    """Test DocumentBuilder with custom font settings."""
    builder = DocumentBuilder(title="Custom", font_name="Arial", font_size=12)
    assert builder.font_name == "Arial"
    assert builder.font_size == 12


def test_document_builder_set_page_size(document_builder: DocumentBuilder) -> None:
    """Test page size is set correctly."""
    section = document_builder.doc.sections[0]
    # Allow for small variations in EMU conversion (within 0.5mm tolerance)
    assert abs(section.page_height - Cm(29.7)) < Cm(0.05)
    assert abs(section.page_width - Cm(21.0)) < Cm(0.05)


def test_document_builder_add_heading(document_builder: DocumentBuilder) -> None:
    """Test adding heading to document."""
    document_builder.add_heading("Heading Level 1", level=1)
    assert len(document_builder.doc.paragraphs) > 0
    assert document_builder.doc.paragraphs[0].text == "Heading Level 1"


def test_document_builder_add_heading_multiple_levels(document_builder: DocumentBuilder) -> None:
    """Test adding headings at different levels."""
    document_builder.add_heading("Title", level=0)
    document_builder.add_heading("Main Heading", level=1)
    document_builder.add_heading("Sub Heading", level=2)

    assert len(document_builder.doc.paragraphs) >= 3


def test_document_builder_add_text(document_builder: DocumentBuilder) -> None:
    """Test adding text paragraph."""
    document_builder.add_text("This is a test paragraph.")

    # Find the paragraph with our text
    paragraphs = [p.text for p in document_builder.doc.paragraphs]
    assert "This is a test paragraph." in paragraphs


def test_document_builder_add_table(document_builder: DocumentBuilder) -> None:
    """Test adding table to document."""
    df = pd.DataFrame({"Name": ["Alice", "Bob"], "Age": [25, 30]})
    column_widths: List[Cm] = [Cm(5), Cm(2)]
    document_builder.add_table(df, column_widths)

    assert len(document_builder.doc.tables) > 0
    table = document_builder.doc.tables[0]
    assert len(table.rows) == 3  # Header + 2 data rows
    assert len(table.columns) == 2


def test_document_builder_add_table_with_custom_font_size(document_builder: DocumentBuilder) -> None:
    """Test adding table with custom font size."""
    df = pd.DataFrame({"Col1": [1, 2], "Col2": [3, 4]})
    column_widths: List[Cm] = [Cm(3), Cm(3)]
    document_builder.add_table(df, column_widths, font_size=11)

    assert len(document_builder.doc.tables) > 0


def test_document_builder_add_table_data(document_builder: DocumentBuilder) -> None:
    """Test add_table_data convenience method."""
    df = pd.DataFrame({"A": [1], "B": [2], "C": [3], "D": [4]})
    document_builder.add_table_data(df)

    assert len(document_builder.doc.tables) > 0


def test_document_builder_add_table_dcr_green(document_builder: DocumentBuilder) -> None:
    """Test add_table_dcr with DCR < 1 (green)."""
    df = pd.DataFrame({"Item": ["Test"], "Value": [100], "DCR": [0.85], "Unit": ["kN"]})
    document_builder.add_table_dcr(df)

    assert len(document_builder.doc.tables) > 0


def test_document_builder_add_table_dcr_red(document_builder: DocumentBuilder) -> None:
    """Test add_table_dcr with DCR >= 1 (red)."""
    df = pd.DataFrame({"Item": ["Test"], "Value": [100], "DCR": [1.2], "Unit": ["kN"]})
    document_builder.add_table_dcr(df)

    assert len(document_builder.doc.tables) > 0


def test_document_builder_add_table_min_max(document_builder: DocumentBuilder) -> None:
    """Test add_table_min_max convenience method."""
    df = pd.DataFrame({"Item": ["Test"], "Min": [1], "Max": [10], "Avg": [5], "Val1": [3], "Val2": [7]})
    document_builder.add_table_min_max(df)

    assert len(document_builder.doc.tables) > 0


def test_document_builder_add_figure(document_builder: DocumentBuilder) -> None:
    """Test adding matplotlib figure to document."""
    # Create a simple figure
    fig, ax = plt.subplots()
    ax.plot([1, 2, 3], [1, 4, 9])
    ax.set_title("Test Plot")

    document_builder.add_figure(fig, width=12)

    # Check that an image was added (images are stored in document parts)
    assert len(document_builder.doc.element.body) > 0


def test_document_builder_add_figure_default_width(document_builder: DocumentBuilder) -> None:
    """Test adding figure with default width."""
    fig, ax = plt.subplots()
    ax.plot([1, 2], [1, 2])

    document_builder.add_figure(fig)  # Uses default width=16

    assert len(document_builder.doc.element.body) > 0


def test_document_builder_add_image(document_builder: DocumentBuilder, tmp_path: Any) -> None:
    """Test adding external image to document."""
    # Create a temporary image file
    fig, ax = plt.subplots()
    ax.plot([1, 2], [1, 2])
    image_path = tmp_path / "test_image.png"
    fig.savefig(str(image_path))
    plt.close(fig)

    document_builder.add_image(str(image_path), width=Cm(10))

    assert len(document_builder.doc.element.body) > 0


def test_document_builder_add_image_no_width(document_builder: DocumentBuilder, tmp_path: Any) -> None:
    """Test adding image without specifying width."""
    # Create a temporary image file
    fig, ax = plt.subplots()
    ax.plot([1, 2], [1, 2])
    image_path = tmp_path / "test_image2.png"
    fig.savefig(str(image_path))
    plt.close(fig)

    document_builder.add_image(str(image_path))

    assert len(document_builder.doc.element.body) > 0


def test_document_builder_save(document_builder: DocumentBuilder, tmp_path: Any) -> None:
    """Test saving document to file."""
    filename = tmp_path / "test_document.docx"
    document_builder.save(str(filename))
    assert filename.exists()


def test_document_builder_full_workflow(tmp_path: Any) -> None:
    """Test a complete workflow with multiple elements."""
    builder = DocumentBuilder(title="Complete Test")

    # Add various elements
    builder.add_heading("Main Title", level=0)
    builder.add_text("Introduction paragraph.")
    builder.add_heading("Section 1", level=1)

    # Use 4-column DataFrame to match add_table_data expectations
    df = pd.DataFrame({"A": [1, 2], "B": [3, 4], "C": [5, 6], "D": [7, 8]})
    builder.add_table_data(df)

    fig, ax = plt.subplots(figsize=(6, 4))
    ax.plot([1, 2, 3], [1, 2, 3])
    builder.add_figure(fig, width=10)

    # Save and verify
    filename = tmp_path / "complete_test.docx"
    builder.save(str(filename))
    assert filename.exists()


if __name__ == "__main__":
    pytest.main()
