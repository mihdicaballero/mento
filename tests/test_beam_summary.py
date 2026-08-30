"""
Comprehensive tests for BeamSummary class in beam_summary.py

This test suite provides 100% coverage of the BeamSummary class,
testing all methods, edge cases, and error conditions.
"""

import pytest
import pandas as pd
import copy
import warnings
from typing import Any, Optional

from docx.oxml.ns import qn
from docx.shared import Emu
from pathlib import Path

from mento import MPa, mm, cm, kN, kNm, m
from mento.beam_summary import BeamSummary
from mento.material import Concrete_ACI_318_19, SteelBar, Concrete_EN_1992_2004
from mento.node import Node
from mento.results import DocumentBuilder, FAIL_MARK, PASS_MARK

# Suppress the specific ACI warning for all tests
pytestmark = pytest.mark.filterwarnings("ignore::UserWarning")

# ============================================================================
# FIXTURES
# ============================================================================


@pytest.fixture
def sample_concrete() -> Concrete_ACI_318_19:
    """Create a sample concrete material."""
    return Concrete_ACI_318_19(name="C25", f_c=25 * MPa)


@pytest.fixture
def sample_steel() -> SteelBar:
    """Create a sample steel material."""
    return SteelBar(name="ADN 420", f_y=420 * MPa)


@pytest.fixture
def sample_input_dataframe() -> pd.DataFrame:
    """
    Create a sample input DataFrame with units and beam data.
    First row contains units, subsequent rows contain beam data.
    """
    data = {
        "Label": ["", "V101", "V102", "V103", "V104"],
        "Comb.": ["", "ELU 1", "ELU 2", "ELU 3", "ELU 4"],
        "b": ["cm", 20, 20, 20, 20],
        "h": ["cm", 50, 50, 50, 50],
        "cc": ["mm", 25, 25, 25, 25],
        "Nx": ["kN", 0, 0, 0, 0],
        "Vz": ["kN", 20, -50, 100, 100],
        "My": ["kNm", 0, -35, 40, 45],
        "ns": ["", 0, 1.0, 1.0, 1.0],
        "dbs": ["mm", 0, 6, 6, 6],
        "sl": ["cm", 0, 20, 20, 20],
        "n1": ["", 2.0, 2, 2.0, 2.0],
        "db1": ["mm", 12, 12, 12, 12],
        "n2": ["", 1.0, 1, 1.0, 0.0],
        "db2": ["mm", 10, 16, 10, 0],
        "n3": ["", 2.0, 0.0, 2.0, 0.0],
        "db3": ["mm", 12, 0, 16, 0],
        "n4": ["", 0, 0.0, 0, 0.0],
        "db4": ["mm", 0, 0, 0, 0],
    }
    return pd.DataFrame(data)


@pytest.fixture
def sample_input_with_nan() -> pd.DataFrame:
    """
    Create a sample input DataFrame with NaN values to test fillna behavior.
    """
    data = {
        "Label": ["", "V101", "V102"],
        "Comb.": ["", "ELU 1", "ELU 2"],
        "b": ["cm", 20, 20],
        "h": ["cm", 50, 50],
        "cc": ["mm", 25, 25],
        "Nx": ["kN", None, 0],
        "Vz": ["kN", 20, None],
        "My": ["kNm", 0, -35],
        "ns": [None, None, 1.0],
        "dbs": ["mm", 0, 6],
        "sl": ["cm", None, 20],
        "n1": ["", 2.0, 2],
        "db1": ["mm", 12, 12],
        "n2": ["", None, 1],
        "db2": ["mm", 10, 16],
        "n3": ["", 2.0, 0.0],
        "db3": ["mm", 12, 0],
        "n4": ["", 0, 0.0],
        "db4": ["mm", 0, 0],
    }
    return pd.DataFrame(data)


@pytest.fixture
def beam_summary(
    sample_concrete: Concrete_ACI_318_19,
    sample_steel: SteelBar,
    sample_input_dataframe: pd.DataFrame,
) -> BeamSummary:
    """Create a BeamSummary instance with sample data."""
    return BeamSummary(
        concrete=sample_concrete,
        steel_bar=sample_steel,
        beam_list=sample_input_dataframe,
    )


# ============================================================================
# INITIALIZATION TESTS
# ============================================================================


def test_beam_summary_initialization(beam_summary: BeamSummary) -> None:
    """Test that BeamSummary initializes correctly."""
    assert isinstance(beam_summary, BeamSummary)
    assert isinstance(beam_summary.concrete, Concrete_ACI_318_19)
    assert isinstance(beam_summary.steel_bar, SteelBar)
    assert isinstance(beam_summary.data, pd.DataFrame)
    assert len(beam_summary.nodes) == 4  # 4 beams in sample data
    assert len(beam_summary.units_row) > 0


def test_beam_summary_with_nan_values(
    sample_concrete: Concrete_ACI_318_19,
    sample_steel: SteelBar,
    sample_input_with_nan: pd.DataFrame,
) -> None:
    """Test that BeamSummary handles NaN values correctly."""
    summary = BeamSummary(
        concrete=sample_concrete,
        steel_bar=sample_steel,
        beam_list=sample_input_with_nan,
    )

    # Check that NaN values were filled with 0
    assert summary.data["Nx"].iloc[0].magnitude == 0
    assert summary.data["Vz"].iloc[1].magnitude == 0
    assert summary.data["sl"].iloc[0].magnitude == 0


def test_beam_summary_with_en_1992_concrete(
    sample_steel: SteelBar,
    sample_input_dataframe: pd.DataFrame,
) -> None:
    """Test BeamSummary with EN 1992 concrete material."""
    concrete_en = Concrete_EN_1992_2004(name="C25/30", f_c=25 * MPa)
    summary = BeamSummary(
        concrete=concrete_en,
        steel_bar=sample_steel,
        beam_list=sample_input_dataframe,
    )
    assert isinstance(summary.concrete, Concrete_EN_1992_2004)


# ============================================================================
# VALIDATE_UNITS TESTS
# ============================================================================


def test_validate_units_valid(beam_summary: BeamSummary) -> None:
    """Test validation with all valid units."""
    valid_units = ["m", "mm", "cm", "inch", "ft", "kN", "kNm", ""]
    # Should not raise any exception
    beam_summary.validate_units(valid_units)


def test_validate_units_invalid(beam_summary: BeamSummary) -> None:
    """Test validation with invalid units."""
    invalid_units = ["m", "mm", "invalid_unit", "kN"]
    with pytest.raises(ValueError, match="Invalid unit 'invalid_unit' detected"):
        beam_summary.validate_units(invalid_units)


def test_validate_units_with_empty_strings(beam_summary: BeamSummary) -> None:
    """Test validation handles empty strings correctly."""
    units = ["", "", "cm", ""]
    # Should not raise any exception
    beam_summary.validate_units(units)


# ============================================================================
# GET_UNIT_VARIABLE TESTS
# ============================================================================


def test_get_unit_variable_valid_units(beam_summary: BeamSummary) -> None:
    """Test getting unit variables for valid units."""
    assert beam_summary.get_unit_variable("mm") == mm
    assert beam_summary.get_unit_variable("cm") == cm
    assert beam_summary.get_unit_variable("m") == m
    assert beam_summary.get_unit_variable("kN") == kN
    assert beam_summary.get_unit_variable("kNm") == kNm
    assert beam_summary.get_unit_variable("MPa") == MPa


def test_get_unit_variable_invalid_unit(beam_summary: BeamSummary) -> None:
    """Test getting unit variable for invalid unit."""
    with pytest.raises(ValueError, match="Unit 'invalid' is not recognized"):
        beam_summary.get_unit_variable("invalid")


# ============================================================================
# CHECK_AND_PROCESS_INPUT TESTS
# ============================================================================


def test_check_and_process_input_creates_data(beam_summary: BeamSummary) -> None:
    """Test that check_and_process_input creates properly formatted data."""
    assert beam_summary.data is not None
    assert len(beam_summary.data) == 4  # 4 data rows
    assert "Label" in beam_summary.data.columns
    assert "My" in beam_summary.data.columns


def test_check_and_process_input_applies_units(beam_summary: BeamSummary) -> None:
    """Test that units are properly applied to columns."""
    # Check that values have proper units
    assert hasattr(beam_summary.data["b"].iloc[0], "magnitude")
    assert hasattr(beam_summary.data["My"].iloc[0], "magnitude")

    # Check units are correct
    assert beam_summary.data["b"].iloc[0].units == cm
    assert beam_summary.data["h"].iloc[0].units == cm
    assert beam_summary.data["My"].iloc[0].units == kNm


# ============================================================================
# CONVERT_TO_NODES TESTS
# ============================================================================


def test_convert_to_nodes_creates_nodes(beam_summary: BeamSummary) -> None:
    """Test that convert_to_nodes creates Node objects."""
    assert len(beam_summary.nodes) == 4
    assert all(isinstance(node, Node) for node in beam_summary.nodes)


def test_convert_to_nodes_with_stirrups(beam_summary: BeamSummary) -> None:
    """Test node creation with stirrups (ns != 0)."""
    # V102, V103, V104 have stirrups
    nodes_with_stirrups = [node for node in beam_summary.nodes[1:]]
    assert len(nodes_with_stirrups) == 3

    # Check that stirrups were set
    for node in nodes_with_stirrups:
        beam = node.section
        assert beam.reinforcement.transverse.n_stirrups != 0


def test_convert_to_nodes_without_stirrups(beam_summary: BeamSummary) -> None:
    """Test node creation without stirrups (ns == 0)."""
    # V101 has no stirrups
    node_without_stirrups = beam_summary.nodes[0]
    beam = node_without_stirrups.section
    assert beam.reinforcement.transverse.n_stirrups == 0


def test_convert_to_nodes_positive_moment(beam_summary: BeamSummary) -> None:
    """Test that positive moments set bottom rebar."""
    # V103 and V104 have positive moments
    for i in [2, 3]:
        node = beam_summary.nodes[i]
        beam = node.section
        # Should have bottom rebar set on the first layer
        assert beam.reinforcement.bottom.layers[0].n != 0


def test_convert_to_nodes_negative_moment(beam_summary: BeamSummary) -> None:
    """Test that negative moments set top rebar."""
    # V102 has negative moment
    node = beam_summary.nodes[1]
    beam = node.section
    # Should have top rebar set on the first layer
    assert beam.reinforcement.top.layers[0].n != 0


def test_convert_to_nodes_zero_moment(beam_summary: BeamSummary) -> None:
    """Test handling of zero moment."""
    # V101 has zero moment (should be treated as positive)
    node = beam_summary.nodes[0]
    beam = node.section
    # Should have bottom rebar set
    assert beam.reinforcement.bottom.layers[0].n != 0


# ============================================================================
# CHECK METHOD TESTS
# ============================================================================


def test_check_method_returns_dataframe(beam_summary: BeamSummary) -> None:
    """Test that check() returns a DataFrame."""
    result = beam_summary.check()
    assert isinstance(result, pd.DataFrame)
    assert len(result) > 0


def test_check_method_capacity_check_false(beam_summary: BeamSummary) -> None:
    """Test check() with capacity_check=False (default)."""
    with warnings.catch_warnings():
        warnings.simplefilter("ignore", UserWarning)
        result = beam_summary.check(capacity_check=False)
    assert isinstance(result, pd.DataFrame)
    # Check that required columns exist
    expected_columns = ["Beam", "b", "h", "As,top", "As,bot", "Av"]
    for col in expected_columns:
        assert col in result.columns


def test_check_method_capacity_check_true(beam_summary: BeamSummary) -> None:
    """Test check() with capacity_check=True."""
    with warnings.catch_warnings():
        warnings.simplefilter("ignore", UserWarning)
        result = beam_summary.check(capacity_check=True)
    assert isinstance(result, pd.DataFrame)
    # Check that result is a DataFrame with expected structure
    assert "Beam" in result.columns
    assert len(result) > 0


def test_check_preserves_forces(beam_summary: BeamSummary) -> None:
    """Test that check() preserves original forces after capacity check."""
    # Store original forces
    original_forces = [[copy.deepcopy(f) for f in node.get_forces_list()] for node in beam_summary.nodes]

    # Run capacity check
    beam_summary.check(capacity_check=True)

    # Verify forces were restored
    for i, node in enumerate(beam_summary.nodes):
        current_forces = node.get_forces_list()
        assert len(current_forces) == len(original_forces[i])


# ============================================================================
# DESIGN METHOD TESTS
# ============================================================================


def test_design_method_returns_dataframe(beam_summary: BeamSummary) -> None:
    """Test that design() returns a DataFrame."""
    result = beam_summary.design()
    assert isinstance(result, pd.DataFrame)
    assert len(result) > 0


def test_design_creates_design_data_attribute(beam_summary: BeamSummary) -> None:
    """Test that design() creates the design_data attribute."""
    beam_summary.design()
    assert hasattr(beam_summary, "design_data")
    assert isinstance(beam_summary.design_data, pd.DataFrame)


def test_design_updates_rebar_columns(beam_summary: BeamSummary) -> None:
    """Test that design() updates rebar columns."""
    result = beam_summary.design()

    # Check that rebar columns exist and are populated
    rebar_columns = ["ns", "dbs", "sl", "n1", "db1", "n2", "db2", "n3", "db3", "n4", "db4"]
    for col in rebar_columns:
        assert col in result.columns


def test_design_handles_positive_moments(beam_summary: BeamSummary) -> None:
    """Test design() properly handles positive moments (bottom rebar)."""
    result = beam_summary.design()

    # V103 and V104 have positive moments - should have bottom rebar designed
    for i in [2, 3]:
        assert result.loc[i, "n1"] > 0  # Should have designed rebar


def test_design_handles_negative_moments(beam_summary: BeamSummary) -> None:
    """Test design() properly handles negative moments (top rebar)."""
    result = beam_summary.design()

    # V102 has negative moment - should have top rebar designed
    assert result.loc[1, "n1"] > 0


# ============================================================================
# SHEAR_RESULTS TESTS
# ============================================================================


def test_shear_results_single_beam(beam_summary: BeamSummary) -> None:
    """Test shear_results() for a single beam."""
    result = beam_summary.shear_results(index=1)
    assert isinstance(result, pd.DataFrame)
    # First row should be units
    assert result.iloc[0, 0] == "" or pd.isna(result.iloc[0, 0])


def test_shear_results_all_beams(beam_summary: BeamSummary) -> None:
    """Test shear_results() for all beams."""
    result = beam_summary.shear_results()
    assert isinstance(result, pd.DataFrame)
    # Should have units row + 4 data rows
    assert len(result) == 5
    # First row should be units
    assert result.iloc[0, 0] == "" or pd.isna(result.iloc[0, 0])


def test_shear_results_index_out_of_range(beam_summary: BeamSummary) -> None:
    """Test shear_results() with out-of-range index."""
    with pytest.raises(IndexError, match="Index .* is out of range"):
        beam_summary.shear_results(index=100)


def test_shear_results_with_capacity_check(beam_summary: BeamSummary) -> None:
    """Test shear_results() with capacity_check=True."""
    result = beam_summary.shear_results(capacity_check=True)
    assert isinstance(result, pd.DataFrame)


def test_shear_results_index_zero_handled(beam_summary: BeamSummary) -> None:
    """Test that index=0 is handled correctly (should use index=0 after max(index-1, 0))."""
    result = beam_summary.shear_results(index=0)
    assert isinstance(result, pd.DataFrame)


# ============================================================================
# FLEXURE_RESULTS TESTS
# ============================================================================


def test_flexure_results_single_beam(beam_summary: BeamSummary) -> None:
    """Test flexure_results() for a single beam."""
    result = beam_summary.flexure_results(index=1)
    assert isinstance(result, pd.DataFrame)
    # First row should be units
    assert result.iloc[0, 0] == "" or pd.isna(result.iloc[0, 0])


def test_flexure_results_all_beams(beam_summary: BeamSummary) -> None:
    """Test flexure_results() for all beams."""
    result = beam_summary.flexure_results()
    assert isinstance(result, pd.DataFrame)
    # Should have units row + 4 data rows
    assert len(result) == 5
    # First row should be units
    assert result.iloc[0, 0] == "" or pd.isna(result.iloc[0, 0])


def test_flexure_results_index_out_of_range(beam_summary: BeamSummary) -> None:
    """Test flexure_results() with out-of-range index."""
    with pytest.raises(IndexError, match="Index .* is out of range"):
        beam_summary.flexure_results(index=100)


def test_flexure_results_with_capacity_check(beam_summary: BeamSummary) -> None:
    """Test flexure_results() with capacity_check=True."""
    result = beam_summary.flexure_results(capacity_check=True)
    assert isinstance(result, pd.DataFrame)


def test_flexure_results_index_zero_handled(beam_summary: BeamSummary) -> None:
    """Test that index=0 is handled correctly."""
    result = beam_summary.flexure_results(index=0)
    assert isinstance(result, pd.DataFrame)


# ============================================================================
# _PROCESS_BEAM_FOR_CHECK TESTS
# ============================================================================


def test_process_beam_for_check_shear(beam_summary: BeamSummary) -> None:
    """Test _process_beam_for_check for shear."""
    node = beam_summary.nodes[0]
    result = beam_summary._process_beam_for_check(node, "shear", False)
    assert isinstance(result, pd.DataFrame)


def test_process_beam_for_check_flexure(beam_summary: BeamSummary) -> None:
    """Test _process_beam_for_check for flexure."""
    node = beam_summary.nodes[0]
    result = beam_summary._process_beam_for_check(node, "flexure", False)
    assert isinstance(result, pd.DataFrame)


def test_process_beam_for_check_invalid_type(beam_summary: BeamSummary) -> None:
    """Test _process_beam_for_check with invalid check_type."""
    node = beam_summary.nodes[0]
    with pytest.raises(ValueError, match="check_type must be either 'shear' or 'flexure'"):
        beam_summary._process_beam_for_check(node, "invalid", False)


def test_process_beam_for_check_capacity_check_restores_forces(
    beam_summary: BeamSummary,
) -> None:
    """Test that _process_beam_for_check restores forces after capacity check."""
    node = beam_summary.nodes[0]
    original_forces = [copy.deepcopy(f) for f in node.get_forces_list()]

    # Run capacity check
    beam_summary._process_beam_for_check(node, "shear", capacity_check=True)

    # Verify forces were restored
    current_forces = node.get_forces_list()
    assert len(current_forces) == len(original_forces)


# ============================================================================
# EXPORT_DESIGN TESTS
# ============================================================================


def test_export_design_without_design_data(beam_summary: BeamSummary) -> None:
    """Test export_design() raises error if design() hasn't been called."""
    with pytest.raises(AttributeError, match="No design data found"):
        beam_summary.export_design("/tmp/test.xlsx")


def test_export_design_creates_file(beam_summary: BeamSummary, tmp_path: Path) -> None:
    """Test export_design() creates an Excel file."""
    # First run design
    beam_summary.design()

    # Export to temp file
    export_path = tmp_path / "test_export.xlsx"
    beam_summary.export_design(str(export_path))

    # Verify file was created
    assert export_path.exists()


def test_export_design_includes_units_row(beam_summary: BeamSummary, tmp_path: Path) -> None:
    """Test export_design() includes units row."""
    # Run design
    beam_summary.design()

    # Export and re-import
    export_path = tmp_path / "test_export.xlsx"
    beam_summary.export_design(str(export_path))

    # Read back the file
    reimported = pd.read_excel(export_path)

    # First row should be units row
    assert len(reimported) == 5  # units + 4 data rows


def test_export_design_converts_magnitudes(beam_summary: BeamSummary, tmp_path: Path) -> None:
    """Test export_design() converts quantities to magnitudes."""
    # Run design
    beam_summary.design()

    # Export
    export_path = tmp_path / "test_export.xlsx"
    beam_summary.export_design(str(export_path))

    # Read back
    reimported = pd.read_excel(export_path)

    # Check that values are numeric (not Quantity objects)
    # Skip first row (units) and first two columns (Label, Comb.)
    for col in reimported.columns[2:]:
        for val in reimported[col].iloc[1:]:  # Skip units row
            assert isinstance(val, (int, float)) or pd.isna(val)


# ============================================================================
# IMPORT_DESIGN TESTS
# ============================================================================


def test_import_design_updates_beam_list(beam_summary: BeamSummary, tmp_path: Path) -> None:
    """Test import_design() updates beam_list."""
    # Export first
    beam_summary.design()
    export_path = tmp_path / "test_import.xlsx"
    beam_summary.export_design(str(export_path))

    # Import
    beam_summary.import_design(str(export_path))

    # Verify beam_list was updated
    assert beam_summary.beam_list is not None
    assert len(beam_summary.beam_list) == 5  # units + 4 beams


def test_import_design_reprocesses_nodes(beam_summary: BeamSummary, tmp_path: Path) -> None:
    """Test import_design() recreates nodes."""
    # Get original node count
    original_node_count = len(beam_summary.nodes)

    # Export and import
    beam_summary.design()
    export_path = tmp_path / "test_import.xlsx"
    beam_summary.export_design(str(export_path))
    beam_summary.import_design(str(export_path))

    # Verify nodes were recreated
    assert len(beam_summary.nodes) == original_node_count


# ============================================================================
# EDGE CASES AND ERROR HANDLING
# ============================================================================


def test_beam_summary_with_minimal_data(
    sample_concrete: Concrete_ACI_318_19,
    sample_steel: SteelBar,
) -> None:
    """Test BeamSummary with minimal valid data."""
    data = {
        "Label": ["", "V1"],
        "Comb.": ["", "ELU 1"],
        "b": ["cm", 20],
        "h": ["cm", 50],
        "cc": ["mm", 25],
        "Nx": ["kN", 0],
        "Vz": ["kN", 20],
        "My": ["kNm", 0],
        "ns": ["", 0],
        "dbs": ["mm", 0],
        "sl": ["cm", 0],
        "n1": ["", 2.0],
        "db1": ["mm", 12],
        "n2": ["", 0.0],
        "db2": ["mm", 0],
        "n3": ["", 0.0],
        "db3": ["mm", 0],
        "n4": ["", 0.0],
        "db4": ["mm", 0],
    }
    df = pd.DataFrame(data)

    summary = BeamSummary(
        concrete=sample_concrete,
        steel_bar=sample_steel,
        beam_list=df,
    )

    assert len(summary.nodes) == 1


def test_beam_summary_all_rebar_layers(
    sample_concrete: Concrete_ACI_318_19,
    sample_steel: SteelBar,
) -> None:
    """Test BeamSummary with all four rebar layers populated."""
    data = {
        "Label": ["", "V1"],
        "Comb.": ["", "ELU 1"],
        "b": ["cm", 20],
        "h": ["cm", 50],
        "cc": ["mm", 25],
        "Nx": ["kN", 0],
        "Vz": ["kN", 20],
        "My": ["kNm", 40],
        "ns": ["", 1],
        "dbs": ["mm", 6],
        "sl": ["cm", 20],
        "n1": ["", 2.0],
        "db1": ["mm", 12],
        "n2": ["", 2.0],
        "db2": ["mm", 10],
        "n3": ["", 2.0],
        "db3": ["mm", 16],
        "n4": ["", 2.0],
        "db4": ["mm", 12],
    }
    df = pd.DataFrame(data)

    summary = BeamSummary(
        concrete=sample_concrete,
        steel_bar=sample_steel,
        beam_list=df,
    )

    beam = summary.nodes[0].section
    # Should have 4 layers of bottom rebar (positive moment)
    bottom = beam.reinforcement.bottom
    assert len(bottom.layers) == 4
    assert [layer.n for layer in bottom.layers] == [2, 2, 2, 2]


def test_beam_summary_with_different_units(
    sample_concrete: Concrete_ACI_318_19,
    sample_steel: SteelBar,
) -> None:
    """Test BeamSummary with different unit combinations."""
    data = {
        "Label": ["", "V1"],
        "Comb.": ["", "ELU 1"],
        "b": ["m", 0.2],
        "h": ["m", 0.5],
        "cc": ["cm", 2.5],
        "Nx": ["kN", 0],
        "Vz": ["kN", 20],
        "My": ["kNm", 40],
        "ns": ["", 1],
        "dbs": ["mm", 6],
        "sl": ["cm", 20],
        "n1": ["", 2.0],
        "db1": ["mm", 12],
        "n2": ["", 0.0],
        "db2": ["mm", 0],
        "n3": ["", 0.0],
        "db3": ["mm", 0],
        "n4": ["", 0.0],
        "db4": ["mm", 0],
    }
    df = pd.DataFrame(data)

    summary = BeamSummary(
        concrete=sample_concrete,
        steel_bar=sample_steel,
        beam_list=df,
    )

    # Units should be properly converted
    assert summary.data["b"].iloc[0].units == m
    assert summary.data["cc"].iloc[0].units == cm


# ============================================================================
# INTEGRATION TESTS
# ============================================================================


def test_full_workflow_check_design_export_import(
    beam_summary: BeamSummary,
    tmp_path: Path,
) -> None:
    """Test complete workflow: check -> design -> export -> import."""
    # 1. Check
    check_result = beam_summary.check()
    assert isinstance(check_result, pd.DataFrame)

    # 2. Design
    design_result = beam_summary.design()
    assert isinstance(design_result, pd.DataFrame)

    # 3. Export
    export_path = tmp_path / "workflow_test.xlsx"
    beam_summary.export_design(str(export_path))
    assert export_path.exists()

    # 4. Import
    beam_summary.import_design(str(export_path))
    assert len(beam_summary.nodes) == 4


def test_check_and_results_consistency(beam_summary: BeamSummary) -> None:
    """Test that check() and separate results methods are consistent."""
    # Run check
    check_df = beam_summary.check()

    # Run separate results
    shear_df = beam_summary.shear_results()
    flexure_df = beam_summary.flexure_results()

    # All should be DataFrames
    assert isinstance(check_df, pd.DataFrame)
    assert isinstance(shear_df, pd.DataFrame)
    assert isinstance(flexure_df, pd.DataFrame)


def test_capacity_check_workflow(beam_summary: BeamSummary) -> None:
    """Test capacity check workflow."""
    # Check with forces
    result_with_forces = beam_summary.check(capacity_check=False)

    # Check capacity
    result_capacity = beam_summary.check(capacity_check=True)

    # Both should be DataFrames but with different content
    assert isinstance(result_with_forces, pd.DataFrame)
    assert isinstance(result_capacity, pd.DataFrame)
    assert len(result_with_forces) == len(result_capacity)


# ============================================================================
# PERFORMANCE AND EDGE CASE TESTS
# ============================================================================


def test_large_number_of_beams(
    sample_concrete: Concrete_ACI_318_19,
    sample_steel: SteelBar,
) -> None:
    """Test BeamSummary with a large number of beams."""
    # Create 20 beams
    n_beams = 20
    data = {
        "Label": [""] + [f"V{i}" for i in range(1, n_beams + 1)],
        "Comb.": [""] + [f"ELU {i}" for i in range(1, n_beams + 1)],
        "b": ["cm"] + [20] * n_beams,
        "h": ["cm"] + [50] * n_beams,
        "cc": ["mm"] + [25] * n_beams,
        "Nx": ["kN"] + [0] * n_beams,
        "Vz": ["kN"] + [20] * n_beams,
        "My": ["kNm"] + [i * 5 for i in range(n_beams)],
        "ns": [""] + [1] * n_beams,
        "dbs": ["mm"] + [6] * n_beams,
        "sl": ["cm"] + [20] * n_beams,
        "n1": [""] + [2] * n_beams,
        "db1": ["mm"] + [12] * n_beams,
        "n2": [""] + [0] * n_beams,
        "db2": ["mm"] + [0] * n_beams,
        "n3": [""] + [0] * n_beams,
        "db3": ["mm"] + [0] * n_beams,
        "n4": [""] + [0] * n_beams,
        "db4": ["mm"] + [0] * n_beams,
    }
    df = pd.DataFrame(data)

    summary = BeamSummary(
        concrete=sample_concrete,
        steel_bar=sample_steel,
        beam_list=df,
    )

    assert len(summary.nodes) == n_beams

    # Should be able to run check on all beams
    result = summary.check()
    assert len(result) == n_beams + 1


def test_empty_label_handling(
    sample_concrete: Concrete_ACI_318_19,
    sample_steel: SteelBar,
) -> None:
    """Test handling of empty labels."""
    data = {
        "Label": ["", "", ""],
        "Comb.": ["", "ELU 1", "ELU 2"],
        "b": ["cm", 20, 20],
        "h": ["cm", 50, 50],
        "cc": ["mm", 25, 25],
        "Nx": ["kN", 0, 0],
        "Vz": ["kN", 20, 30],
        "My": ["kNm", 0, 10],
        "ns": ["", 0, 1],
        "dbs": ["mm", 0, 6],
        "sl": ["cm", 0, 20],
        "n1": ["", 2, 2],
        "db1": ["mm", 12, 12],
        "n2": ["", 0, 0],
        "db2": ["mm", 0, 0],
        "n3": ["", 0, 0],
        "db3": ["mm", 0, 0],
        "n4": ["", 0, 0],
        "db4": ["mm", 0, 0],
    }
    df = pd.DataFrame(data)

    summary = BeamSummary(
        concrete=sample_concrete,
        steel_bar=sample_steel,
        beam_list=df,
    )

    assert len(summary.nodes) == 2


# ============================================================================
# PROPERTY AND ATTRIBUTE TESTS
# ============================================================================


def test_beam_summary_attributes_set(beam_summary: BeamSummary) -> None:
    """Test that all expected attributes are set on initialization."""
    assert hasattr(beam_summary, "concrete")
    assert hasattr(beam_summary, "steel_bar")
    assert hasattr(beam_summary, "beam_list")
    assert hasattr(beam_summary, "units_row")
    assert hasattr(beam_summary, "data")
    assert hasattr(beam_summary, "nodes")


def test_units_row_populated(beam_summary: BeamSummary) -> None:
    """Test that units_row is properly populated."""
    assert len(beam_summary.units_row) > 0
    # First column should be empty (for Label)
    assert beam_summary.units_row[0] == ""
    # Other columns should have units or be empty
    assert all(isinstance(unit, str) for unit in beam_summary.units_row)


def test_nodes_have_correct_structure(beam_summary: BeamSummary) -> None:
    """Test that nodes have the correct structure."""
    for node in beam_summary.nodes:
        assert hasattr(node, "section")
        assert hasattr(node, "forces")
        assert hasattr(node, "get_forces_list")


def test_check_with_en1992_concrete(
    sample_steel: SteelBar,
    sample_input_dataframe: pd.DataFrame,
) -> None:
    """Test check() with EN 1992 concrete includes code-specific columns."""
    # Create EN 1992 concrete
    concrete_en = Concrete_EN_1992_2004(name="C25/30", f_c=25 * MPa)

    summary = BeamSummary(
        concrete=concrete_en,
        steel_bar=sample_steel,
        beam_list=sample_input_dataframe,
    )

    result = summary.check(capacity_check=False)

    # Check that EN 1992 specific columns exist
    assert "MRd,top" in result.columns
    assert "MRd,bot" in result.columns
    assert "VRd" in result.columns

    # Verify data types (should have numeric values)
    assert isinstance(result, pd.DataFrame)
    assert len(result) > 0


def test_design_with_en1992_concrete(
    sample_steel: SteelBar,
    sample_input_dataframe: pd.DataFrame,
) -> None:
    """design() must work for EN 1992 beams, not only ACI ones.

    Regression test: the flexural design driver is shared between codes, so the
    EN path now fills ``flexure_design_results_bot``/``_top``, which is what
    ``BeamSummary.design`` reads to write the rebar columns back. Before the
    two codes shared a driver, the EN branch never set them and design() failed.
    """
    concrete_en = Concrete_EN_1992_2004(name="C25/30", f_c=25 * MPa)

    summary = BeamSummary(
        concrete=concrete_en,
        steel_bar=sample_steel,
        beam_list=sample_input_dataframe,
    )

    result = summary.design()

    assert isinstance(result, pd.DataFrame)
    assert len(result) > 0
    # Every designed beam ends up with at least one layer of longitudinal bars.
    assert (result["n1"].astype(int) > 0).all()
    assert (result["db1"].apply(lambda q: q.magnitude) > 0).all()


def test_check_with_en1992_capacity_check(
    sample_steel: SteelBar,
    sample_input_dataframe: pd.DataFrame,
) -> None:
    """Test capacity check with EN 1992 concrete."""
    concrete_en = Concrete_EN_1992_2004(name="C25/30", f_c=25 * MPa)

    summary = BeamSummary(
        concrete=concrete_en,
        steel_bar=sample_steel,
        beam_list=sample_input_dataframe,
    )

    flexure_result = summary.flexure_results(capacity_check=True)
    shear_result = summary.shear_results(capacity_check=True)

    # Should return DataFrame with EN 1992 columns
    assert isinstance(flexure_result, pd.DataFrame)
    assert "MRd,top" in flexure_result.columns
    assert "MRd,bot" in flexure_result.columns
    assert "VRd" in shear_result.columns


def test_en1992_vs_aci_column_differences(
    sample_steel: SteelBar,
    sample_input_dataframe: pd.DataFrame,
) -> None:
    """Test that EN 1992 and ACI produce different column sets."""
    # ACI concrete
    concrete_aci = Concrete_ACI_318_19(name="C25", f_c=25 * MPa)
    summary_aci = BeamSummary(
        concrete=concrete_aci,
        steel_bar=sample_steel,
        beam_list=sample_input_dataframe,
    )
    result_aci = summary_aci.flexure_results(capacity_check=True)

    # EN 1992 concrete
    concrete_en = Concrete_EN_1992_2004(name="C25/30", f_c=25 * MPa)
    summary_en = BeamSummary(
        concrete=concrete_en,
        steel_bar=sample_steel,
        beam_list=sample_input_dataframe,
    )
    result_en = summary_en.flexure_results(capacity_check=True)

    # EN 1992 should have MRd columns, ACI should have ØMn columns
    assert "MRd,top" in result_en.columns
    assert "MRd,top" not in result_aci.columns
    assert "ØMn,top" in result_aci.columns
    assert "ØMn,top" not in result_en.columns


def test_check_capacity_true_with_en1992_concrete(
    sample_steel: SteelBar,
    sample_input_dataframe: pd.DataFrame,
) -> None:
    """Test check(capacity_check=True) with EN 1992 includes code-specific capacity columns."""
    concrete_en = Concrete_EN_1992_2004(name="C25/30", f_c=25 * MPa)
    summary = BeamSummary(
        concrete=concrete_en,
        steel_bar=sample_steel,
        beam_list=sample_input_dataframe,
    )
    result = summary.check(capacity_check=True)
    assert isinstance(result, pd.DataFrame)
    assert "MRd,top" in result.columns
    assert "MRd,bot" in result.columns
    assert "VRd" in result.columns


def test_results_detailed_doc_invalid_index(beam_summary: BeamSummary) -> None:
    """Test results_detailed_doc() raises IndexError for out-of-range index."""
    with pytest.raises(IndexError):
        beam_summary.results_detailed_doc(index=0)
    with pytest.raises(IndexError):
        beam_summary.results_detailed_doc(index=100)


def test_results_detailed_doc_aci(
    beam_summary: BeamSummary,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Test results_detailed_doc() runs without error for ACI concrete (non-EN 1992 shear branch)."""
    monkeypatch.setattr(DocumentBuilder, "save", lambda *_: None)
    beam_summary.results_detailed_doc(index=1)


def test_results_detailed_doc_en1992(
    sample_steel: SteelBar,
    sample_input_dataframe: pd.DataFrame,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Test results_detailed_doc() runs without error for EN 1992 concrete (EN 1992 shear branch)."""
    concrete_en = Concrete_EN_1992_2004(name="C25/30", f_c=25 * MPa)
    summary = BeamSummary(
        concrete=concrete_en,
        steel_bar=sample_steel,
        beam_list=sample_input_dataframe,
    )
    monkeypatch.setattr(DocumentBuilder, "save", lambda *_: None)
    summary.results_detailed_doc(index=1)


def test_deprecated_summary_module_still_exports_beam_summary():
    """mento.summary is the pre-rename import path and must keep working."""
    import importlib
    import sys

    # Drop the module so the shim body — and its warning — runs again.
    sys.modules.pop("mento.summary", None)

    with warnings.catch_warnings(record=True) as caught:
        warnings.simplefilter("always")
        legacy = importlib.import_module("mento.summary")

    assert legacy.BeamSummary is BeamSummary
    assert any(issubclass(w.category, DeprecationWarning) for w in caught)
    assert any("beam_summary" in str(w.message) for w in caught)


# ---------------------------------------------------------------------------
# The Word "Design Check Summary" table
# ---------------------------------------------------------------------------


def _built_document(summary: BeamSummary, monkeypatch: pytest.MonkeyPatch) -> Any:
    """Build the Word summary and hand back the document instead of saving it."""
    captured: list = []
    monkeypatch.setattr(DocumentBuilder, "save", lambda self, *_: captured.append(self.doc))
    summary.results_detailed_doc(index=1)
    return captured[0]


def _rows(table: Any) -> list:
    return [[cell.text for cell in row.cells] for row in table.rows]


def _cell_fill(cell: Any) -> Optional[str]:
    """The shading colour of a table cell, or None if it has none."""
    tc_pr = cell._element.find(qn("w:tcPr"))
    if tc_pr is None:
        return None
    shd = tc_pr.find(qn("w:shd"))
    return None if shd is None else shd.get(qn("w:fill"))


def test_check_summary_table_carries_only_the_reporting_columns(
    beam_summary: BeamSummary, monkeypatch: pytest.MonkeyPatch
) -> None:
    """The closing table answers "did every beam pass, and on what".

    The required-area and capacity columns are left out on purpose: the flexure
    and shear tables above already report them in full, and repeating them here
    pushed the table off the page.
    """
    table = _built_document(beam_summary, monkeypatch).tables[-1]

    assert _rows(table)[0] == [
        "Beam",
        "b",
        "h",
        "As,top",
        "As,bot",
        "Av",
        "Mu",
        "Vu",
        "Nu",
        "DCRb,top",
        "DCRb,bot",
        "DCRv",
        "Status",
    ]
    # The units row survives the column selection.
    assert _rows(table)[1][:3] == ["", "cm", "cm"]
    assert _rows(table)[1][6:9] == ["kNm", "kN", "kN"]


def test_check_summary_table_uses_the_design_code_names_for_the_demands(
    sample_steel: SteelBar, sample_input_dataframe: pd.DataFrame, monkeypatch: pytest.MonkeyPatch
) -> None:
    """EN reports M_Ed / V_Ed / N_Ed where ACI reports Mu / Vu / Nu."""
    summary = BeamSummary(
        concrete=Concrete_EN_1992_2004(name="C25/30", f_c=25 * MPa),
        steel_bar=sample_steel,
        beam_list=sample_input_dataframe,
    )
    header = _rows(_built_document(summary, monkeypatch).tables[-1])[0]

    assert header[6:9] == ["MEd", "VEd", "NEd"]


def test_the_status_column_is_shaded_green_when_it_passes(
    beam_summary: BeamSummary, monkeypatch: pytest.MonkeyPatch
) -> None:
    """A reader scans the summary for red rather than reading every tick."""
    table = _built_document(beam_summary, monkeypatch).tables[-1]
    status_idx = len(table.columns) - 1

    verdicts = [(row.cells[0].text, row.cells[status_idx]) for row in table.rows]
    header_label, header_cell = verdicts[0]
    units_label, units_cell = verdicts[1]

    # Neither the header nor the units row carries a verdict, so neither is shaded.
    assert header_cell.text == "Status"
    assert _cell_fill(header_cell) is None
    assert units_cell.text == ""
    assert _cell_fill(units_cell) is None

    beams = verdicts[2:]
    assert beams, "the fixture produced no beam rows"
    for label, cell in beams:
        expected = "C6EFCE" if cell.text == PASS_MARK else "FFC7CE"
        assert _cell_fill(cell) == expected, f"{label} verdict {cell.text!r} was not shaded"


def test_a_beam_that_fails_is_shaded_red(
    sample_concrete: Concrete_ACI_318_19, sample_steel: SteelBar, monkeypatch: pytest.MonkeyPatch
) -> None:
    """The colouring has to distinguish, not just decorate."""
    # The second beam is deliberately far too small for its demand.
    beam_list = pd.DataFrame(
        {
            "Label": ["", "passes", "fails"],
            "Comb.": ["", "ELU 1", "ELU 2"],
            "b": ["cm", 25, 20],
            "h": ["cm", 50, 30],
            "cc": ["mm", 25, 25],
            "Nx": ["kN", 0, 0],
            "Vz": ["kN", 40, 250],
            "My": ["kNm", 40, 200],
            "ns": ["", 1, 1],
            "dbs": ["mm", 8, 6],
            "sl": ["cm", 20, 25],
            "n1": ["", 3, 2],
            "db1": ["mm", 16, 8],
            "n2": ["", 0, 0],
            "db2": ["mm", 0, 0],
            "n3": ["", 0, 0],
            "db3": ["mm", 0, 0],
            "n4": ["", 0, 0],
            "db4": ["mm", 0, 0],
        }
    )
    summary = BeamSummary(concrete=sample_concrete, steel_bar=sample_steel, beam_list=beam_list)
    summary.check()
    table = _built_document(summary, monkeypatch).tables[-1]
    status_idx = len(table.columns) - 1

    by_label = {row.cells[0].text: row.cells[status_idx] for row in table.rows}
    assert by_label["passes"].text == PASS_MARK
    assert _cell_fill(by_label["passes"]) == "C6EFCE"
    assert by_label["fails"].text == FAIL_MARK
    assert _cell_fill(by_label["fails"]) == "FFC7CE"


def test_the_all_beam_tables_are_set_a_point_smaller(
    beam_summary: BeamSummary, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Flexure, shear and the closing check table are much wider than the text."""
    doc = _built_document(beam_summary, monkeypatch)

    for table in doc.tables[-3:]:
        sizes = {
            run.font.size.pt
            for row in table.rows
            for cell in row.cells
            for paragraph in cell.paragraphs
            for run in paragraph.runs
            if run.font.size is not None
        }
        assert sizes == {8.0}, f"expected 8 pt throughout, found {sizes}"


def _table_width_cm(table: Any) -> float:
    return sum(Emu(cell.width).cm for cell in table.rows[0].cells if cell.width is not None)


def _usable_width_cm(doc: Any) -> float:
    section = doc.sections[0]
    return Emu(section.page_width - section.left_margin - section.right_margin).cm


def test_no_table_in_the_report_runs_past_the_page(beam_summary: BeamSummary, monkeypatch: pytest.MonkeyPatch) -> None:
    """Widths are authoritative now, so one that overruns would clip."""
    doc = _built_document(beam_summary, monkeypatch)
    usable = _usable_width_cm(doc)

    for i, table in enumerate(doc.tables):
        assert _table_width_cm(table) <= usable + 0.01, f"table {i} is wider than the text column"


def test_the_per_beam_tables_do_not_span_the_line(beam_summary: BeamSummary, monkeypatch: pytest.MonkeyPatch) -> None:
    """A table describing one beam is sized to its text, not to the page.

    Only the all-beams summaries at the end are meant to span the line; the
    detail tables above them read as a list, and stretching a four-column list
    across the full width leaves the value adrift from its label.
    """
    doc = _built_document(beam_summary, monkeypatch)
    usable = _usable_width_cm(doc)

    # Everything before the "Summary - All Beams" tables describes one beam.
    detail_tables = doc.tables[:-4]
    assert detail_tables, "the report produced no detail tables"
    for i, table in enumerate(detail_tables):
        assert _table_width_cm(table) < usable - 1.0, f"detail table {i} spans the line"


def test_table_widths_are_honoured_rather_than_stored(
    beam_summary: BeamSummary, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Word ignores cell widths unless autofit is off and the layout is fixed.

    Without both, every table is stretched to the text width and the widths
    each caller chose are stored in the file but never used.
    """
    for table in _built_document(beam_summary, monkeypatch).tables:
        assert table.autofit is False
        layout = table._tbl.find(qn("w:tblPr")).find(qn("w:tblLayout"))
        assert layout is not None and layout.get(qn("w:type")) == "fixed"


def test_limit_check_verdicts_are_shaded_like_the_summary(
    beam_summary: BeamSummary, monkeypatch: pytest.MonkeyPatch
) -> None:
    """A limit check is a pass/fail statement too, so it gets the same colours."""
    doc = _built_document(beam_summary, monkeypatch)

    limit_tables = [t for t in doc.tables if t.rows[0].cells[-1].text == "Ok?"]
    assert limit_tables, "the report produced no limit-check tables"

    verdicts = []
    for table in limit_tables:
        idx = len(table.columns) - 1
        assert _cell_fill(table.rows[0].cells[idx]) is None  # the header is not a verdict
        for row in table.rows[1:]:
            cell = row.cells[idx]
            expected = "C6EFCE" if cell.text == PASS_MARK else "FFC7CE"
            assert _cell_fill(cell) == expected, f"{row.cells[0].text!r} was not shaded"
            verdicts.append(cell.text)

    assert verdicts, "the limit-check tables carried no verdicts"


def test_a_failed_limit_check_is_shaded_red(
    sample_concrete: Concrete_ACI_318_19, sample_steel: SteelBar, monkeypatch: pytest.MonkeyPatch
) -> None:
    """The colouring has to distinguish, not just decorate.

    A 6 mm stirrup is below the 10 mm ACI 318-19 detailing minimum, so the
    "Minimum rebar diameter" limit fails while the rest of the table passes.
    """
    beam_list = pd.DataFrame(
        {
            "Label": ["", "thin-stirrup"],
            "Comb.": ["", "ELU 1"],
            "b": ["cm", 25],
            "h": ["cm", 50],
            "cc": ["mm", 25],
            "Nx": ["kN", 0],
            "Vz": ["kN", 72],
            "My": ["kNm", 90],
            "ns": ["", 1],
            "dbs": ["mm", 6],
            "sl": ["cm", 20],
            "n1": ["", 3],
            "db1": ["mm", 16],
            "n2": ["", 0],
            "db2": ["mm", 0],
            "n3": ["", 0],
            "db3": ["mm", 0],
            "n4": ["", 0],
            "db4": ["mm", 0],
        }
    )
    summary = BeamSummary(concrete=sample_concrete, steel_bar=sample_steel, beam_list=beam_list)
    summary.check()
    doc = _built_document(summary, monkeypatch)

    limit_tables = [t for t in doc.tables if t.rows[0].cells[-1].text == "Ok?"]
    by_check = {
        row.cells[0].text: row.cells[len(table.columns) - 1] for table in limit_tables for row in table.rows[1:]
    }

    failed = by_check["Minimum rebar diameter"]
    assert failed.text == FAIL_MARK
    assert _cell_fill(failed) == "FFC7CE"
    # And a passing check in the same table is still green.
    passed = by_check["Stirrup spacing along length"]
    assert passed.text == PASS_MARK
    assert _cell_fill(passed) == "C6EFCE"


def test_the_shear_summary_drops_the_redundant_capacity_tick(
    beam_summary: BeamSummary, monkeypatch: pytest.MonkeyPatch
) -> None:
    """The DCR column beside it already says whether the section is enough."""
    doc = _built_document(beam_summary, monkeypatch)
    # The shear summary is the one carrying Av,min; the third from the end.
    header = [cell.text for cell in doc.tables[-2].rows[0].cells]

    assert "Av,min" in header, f"expected the shear summary, got {header}"
    assert "Vu≤ØVn" not in header
    assert "Vu≤ØVmax" in header, "the maximum-capacity check is still worth reporting"
    assert "DCR" in header
