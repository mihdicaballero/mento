import warnings

import pandas as pd
from typing import Optional, List, Any
from tabulate import tabulate
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
import seaborn as sns
import matplotlib.pyplot as plt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from io import BytesIO
from pandas.io.formats.style import Styler

from mento.i18n import DEFAULT_LANGUAGE, translate, translate_dataframe, translate_table

#: The verdict marks a summary's Status column carries. Named because the Word
#: builder shades that column by matching them: a literal on one side and a
#: different literal on the other would simply stop colouring, silently.
PASS_MARK = "✅"
FAIL_MARK = "❌"

#: What the pass/fail column is called. The limit-check tables already asked
#: "Ok?", and the summaries agreeing with them keeps the column narrow.
VERDICT_COLUMN = "Ok?"

#: Decimals for the reported numbers. Forces and moments are quoted in kN and
#: kN*m, where a second decimal is below what any of these calculations can
#: claim and only widens the column; a DCR is a ratio around 1, where two
#: decimals is the resolution a reader acts on.
FORCE_DECIMALS = 1
DCR_DECIMALS = 2

#: The units a force or a moment is quoted in. A value in one of these is
#: rounded to `FORCE_DECIMALS` for display.
FORCE_DISPLAY_UNITS = {"kN", "kNm", "kN*m", "kN·m", "kip", "kip*ft", "kip·ft"}


def _rounded(value: Any, decimals: int) -> Any:
    """Round a float, and leave anything else -- a label, a tick -- alone."""
    return round(value, decimals) if isinstance(value, float) else value


def round_for_display(df: pd.DataFrame) -> pd.DataFrame:
    """A copy of ``df`` with forces at one decimal and DCRs at two.

    Display only, and deliberately so. The frames themselves keep the
    precision they were built with, because they are also the programmatic
    result and the validation suite compares them against Calcpad and ETABS
    references at more decimals than a report has room to show.

    Two shapes reach a Word table: a detail table, which carries the unit
    beside each value, and an all-rows summary, which carries the units on its
    first row. The unit is what decides, so a column of areas keeps its second
    decimal either way.
    """
    out = df.copy()

    if "Unit" in out.columns and "Value" in out.columns:
        variables = out["Variable"].astype(str) if "Variable" in out.columns else pd.Series([""] * len(out))
        for column in ("Value", "Min.", "Max."):
            if column not in out.columns:
                continue
            out[column] = [
                _rounded(value, DCR_DECIMALS)
                if variable == "DCR"
                else _rounded(value, FORCE_DECIMALS)
                if str(unit) in FORCE_DISPLAY_UNITS
                else value
                for value, unit, variable in zip(out[column], out["Unit"], variables)
            ]
        return out

    # A summary always carries its units on the first row; a frame without one
    # is a bug in whoever built it, and should say so rather than be skipped.
    units = out.iloc[0]
    for column in out.columns:
        if str(column).startswith("DCR"):
            decimals = DCR_DECIMALS
        elif str(units[column]) in FORCE_DISPLAY_UNITS:
            decimals = FORCE_DECIMALS
        else:
            continue
        out[column] = [_rounded(value, decimals) for value in out[column]]
    return out


#: Column widths for a per-element detail table: a description, then a symbol,
#: a value and a unit. Sized to the longest description these tables carry
#: ("Depth of equivalent strength block ratio", 40 characters) rather than to
#: the page, so a table that describes one beam does not span the line the way
#: the all-beams summaries do. With the fixed layout `add_table` sets, anything
#: longer wraps rather than being clipped.
DETAIL_TABLE_WIDTHS = [Cm(6.1), Cm(2.2), Cm(3.0), Cm(1.4)]

#: The same, for a limit check: description, unit, value, minimum, maximum and
#: the verdict.
LIMIT_TABLE_WIDTHS = [Cm(5.5), Cm(1.5), Cm(1.5), Cm(1.5), Cm(1.5), Cm(1.2)]

CUSTOM_COLORS = {
    "blue": "#1f77b4",  # Default Matplotlib blue
    "red": "#d62728",  # Default Matplotlib red
    "dark_gray": "#323232",
    "light_gray": "#e3e3e3",
    "dark_blue": "#073165",
}


# Default Matplotlib and Seaborn settings
def configure_plot_settings() -> None:
    """
    Configures global settings for Matplotlib and Seaborn plots.
    """
    # Basic style for the plot
    sns.set_theme(style="whitegrid")
    sns.set_context("paper", rc={"lines.linewidth": 1.8})
    sns.set_style(rc={"axes.facecolor": "#F8F8F8"})

    # Matplotlib specific settings
    plt.rcParams.update(
        {
            "text.usetex": True,
            "font.family": "serif",
            "font.serif": ["Lato"],
            "axes.titlesize": 12,
            "axes.labelsize": 12,
            "xtick.labelsize": 12,
            "ytick.labelsize": 12,
            "legend.fontsize": 12,
        }
    )


class Formatter:
    """
    Class to stlye the results in a Markdown display and a DataFrame output.
    """

    def __init__(self) -> None:
        # Define colors in the constructor
        self.green = "#439b00"
        self.red = "#d43e36"
        self.yellow = "#efc200"
        self.mid_value = 0.95
        self.max_value = 1

    def DCR(self, DCR: float) -> str:
        # Determine color based on DCR value
        if self.mid_value > DCR:
            color = self.green
        elif self.mid_value <= DCR <= self.max_value:
            color = self.yellow
        else:
            color = self.red

        return f"$\\color{{{color}}}{{\\text{{DCR}}={round(DCR, 2)}}}$"

    def DCR_value(self, DCR: float) -> str:
        # Determine color based on DCR value
        if self.mid_value > DCR:
            color = self.green
        elif self.mid_value <= DCR <= 1:
            color = self.yellow
        else:
            color = self.red

        return f"$\\color{{{color}}}{{{round(DCR, 2)}}}$"

    def is_lower(self, value1: float, value2: float) -> str:
        # Compare two values and return the appropriate formatted output
        if value1 < value2:
            return r"$\color{" + self.green + r"}{\, \checkmark}$"  # Green checkmark
        else:
            return r"$\color{" + self.red + r"}{\, \times}$"  # Red cross

    def is_greater(self, value1: float, value2: float) -> str:
        # Compare two values and return the appropriate formatted output
        if value1 > value2:
            return r"$\color{" + self.green + r"}{\, \checkmark}$"  # Green checkmark
        else:
            return r"$\color{" + self.red + r"}{\, \times}$"  # Red cross

    # Formatting functions outside the class

    def DCR_value_df(self, DCR: Optional[float]) -> str:
        # If DCR is None or not a number, return an empty string
        if DCR is None or not isinstance(DCR, (int, float)):
            return ""

        # Determine the color based on DCR value
        if self.mid_value > DCR:
            color = self.green
        elif self.mid_value <= DCR <= self.max_value:  # Use self.max_value here for consistency
            color = self.yellow
        else:
            color = self.red

        return f"color: {color}"  # This is a CSS style string

    def apply_DCR_style(self, value: float) -> str:
        formatted_value = round(value, 2) if isinstance(value, (int, float)) else value
        return self.DCR_value_df(formatted_value)

    def color_DCR_df(self, df: pd.DataFrame, DCR_columns: list) -> Styler:  # Use Styler directly now
        """
        Apply color styling to specified DCR-related columns in the DataFrame.

        :param df: DataFrame with DCR values to color.
        :param DCR_columns: List of column names to apply the DCR styling to.
        :return: A styled DataFrame with colored DCR values.
        """
        return df.style.map(self.apply_DCR_style, subset=DCR_columns).format(precision=2)


class TablePrinter:
    """
    A class for printing tables with customizable formatting options.

    Attributes
    ----------
    title : Optional[str]
        Optional title displayed above the printed table.

    Methods
    -------
    print_table_data(data: List[List[Any]], headers: List[str], tablefmt: str = "fancygrid", numalign: str = "right")
      -> None
        Prints table data with customizable formatting.

    print_table_min_max(data: List[List[Any]], headers: List[str], tablefmt: str = "fancygrid", numalign: str = "right")
      -> None
        Prints table data with column alignment for minimum and maximum values.
    """

    def __init__(self, title: Optional[str] = None, language: str = DEFAULT_LANGUAGE) -> None:
        """
        Initializes the TablePrinter with an optional title.

        Parameters
        ----------
        title : Optional[str], default=None
            Title to be displayed above the table, if provided.

        language : str, default="en"
            Language the table is rendered in. Headers and row labels without a
            translation are printed in English.
        """
        self.language = language
        self.title = translate(title, language) if title else title

    def print_table_data(
        self,
        data: dict[str, list[Any]],
        headers: str,
        tablefmt: str = "fancygrid",
        numalign: str = "right",
    ) -> None:
        """
        Prints table data with customizable formatting options.

        Parameters
        ----------
        data : List[List[Any]]
            The data to be printed in table format, where each inner list represents a row.

        headers : List[str]
            Column headers for the table.

        tablefmt : str, default="fancygrid"
            Table formatting style supported by tabulate (e.g., "plain", "grid", "fancygrid").

        numalign : str, default="right"
            Number alignment in the table. Common options are "right", "center", or "left".

        Returns
        -------
        table:  Return the formatted table string
        """
        # if self.title:
        #     print(f"======= {self.title} ======= \n")

        colalign = ("left", "center", "right", "left")
        table = tabulate(
            translate_table(data, self.language),
            headers=headers,
            tablefmt=tablefmt,
            numalign=numalign,
            colalign=colalign,
        )
        print(table, "\n")
        return None

    def print_table_min_max(
        self,
        data: dict[str, list[Any]],
        headers: str,
        tablefmt: str = "fancygrid",
        numalign: str = "right",
    ) -> None:
        """
        Prints table data with column alignment for minimum and maximum values.

        Parameters
        ----------
        data : List[List[Any]]
            The data to be printed in table format, where each inner list represents a row.

        headers : List[str]
            Column headers for the table.

        tablefmt : str, default="fancygrid"
            Table formatting style supported by tabulate (e.g., "plain", "grid", "fancygrid").

        numalign : str, default="right"
            Number alignment in the table. Common options are "right", "center", or "left".

        Returns
        -------
        None
        """
        if self.title:
            print(f"======= {self.title} ======= \n")

        colalign = ("left", "center", "center", "center", "center", "center", "left")
        table = tabulate(
            translate_table(data, self.language),
            headers=headers,
            tablefmt=tablefmt,
            numalign=numalign,
            colalign=colalign,
        )
        print(table, "\n")


class DocumentBuilder:
    """
    A class to build and style a Word document, including adding headings and tables from data frames.

    Attributes
    ----------
    title : str
        Title of the document, displayed at the beginning.

    font_name : str
        Default font name for the document text.

    font_size : int
        Default font size for the document text.

    doc : Document
        The Document object representing the Word document being built.

    Methods
    -------
    set_document_style() -> None
        Configures the document's default style with the specified font name and size.

    add_heading(text: str, level: int) -> None
        Adds a heading to the document at the specified level.

    set_col_widths(table: 'docx.table.Table', column_widths: List[Cm]) -> None
        Sets the width for each column in a given table.

    add_table(df: pd.DataFrame, column_widths: List[Cm]) -> None
        Adds a table to the document, with data from a DataFrame and custom column widths.

    save(filename: str) -> None
        Saves the document to a specified filename.
    """

    def __init__(
        self,
        title: str,
        font_name: str = "Lato",
        font_size: int = 9,
        language: str = DEFAULT_LANGUAGE,
    ) -> None:
        """
        Initializes the DocumentBuilder with a title, font name, and font size.

        Parameters
        ----------
        title : str
            Title to be displayed in the document.

        font_name : str, default='Lato'
            Font name to be used for the document text.

        font_size : int, default=9
            Font size for the document text.

        language : str, default="en"
            Language the document is written in. Headings, table headers and row
            labels without a translation are written in English.
        """
        self.doc = Document()
        self.language = language
        self.title = translate(title, language)
        self.font_name = font_name
        self.font_size = font_size
        self.set_document_style()
        self.set_page_size()
        self.set_margins()

    def set_document_style(self) -> None:
        """
        Sets the default style of the document, applying the font name and size.

        Returns
        -------
        None
        """
        # Normal text style
        style = self.doc.styles["Normal"]
        style.font.name = self.font_name
        style.font.size = Pt(self.font_size)

    def set_page_size(self) -> None:
        """
        Sets the page size to A4 by default.

        Returns
        -------
        None
        """
        section = self.doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)

    def set_margins(self, top: float = 3.5, bottom: float = 2.5, left: float = 1.6, right: float = 1.6) -> None:
        """
        Set page margins in centimeters.

        Parameters
        ----------
        top, bottom, left, right : float
            Margins in centimeters. Default is 1 cm for all.
        """
        for section in self.doc.sections:
            section.top_margin = Cm(top)
            section.bottom_margin = Cm(bottom)
            section.left_margin = Cm(left)
            section.right_margin = Cm(right)

    def add_heading(self, text: str, level: int, font_size: Optional[float] = 10, **fields: Any) -> None:
        """
        Adds a heading to the document at the specified level.

        Parameters
        ----------
        text : str
            The text for the heading, in English. It is looked up in the
            document's language catalog before being written.
        level : int
            The heading level (e.g., 0 for title, 1 for main headings, etc.).
        font_size : float, optional
            Font size in points. Default is 10 pt.
        **fields
            Values for the ``{placeholders}`` of ``text``, filled in after the
            translation so both languages can order them differently.

        Returns
        -------
        None
        """
        heading = self.doc.add_heading(translate(text, self.language, **fields), level=level)
        heading.paragraph_format.space_before = Pt(0)

        # Set font size for all runs in heading
        for run in heading.runs:
            run.font.size = Pt(font_size)
            run.font.name = self.font_name

    def add_text(self, text: str, **fields: Any) -> None:
        """Adds a paragraph to the document, translated into its language.

        ``text`` is the English string; ``fields`` fill its ``{placeholders}``.
        """
        self.doc.add_paragraph(translate(text, self.language, **fields))

    def set_col_widths(self, table: Any, column_widths: List[Cm]) -> None:
        """
        Sets the width of each column in the table.

        Parameters
        ----------
        table : docx.table.Table
            The table object whose column widths need to be set.

        column_widths : List[Cm]
            List of widths for each column, in centimeters.

        Returns
        -------
        None
        """
        n_columns = len(table.columns)
        widths = list(column_widths[:n_columns])
        # A caller that gives fewer widths than there are columns used to leave
        # the rest at Word's default, which is what made the flexure limit
        # checks wider than the page. Missing ones repeat the last width given:
        # a trailing run of similar columns is the usual shape, and a fixed
        # fallback made the narrowest columns of a wide table the widest.
        if len(widths) < n_columns:
            # Padding keeps the table renderable, but a short list is a caller
            # bug: it is what made the last columns of Beam Data take a
            # fallback width and squeezed the rest. Say so rather than let the
            # layout be the only evidence.
            warnings.warn(
                f"table has {n_columns} columns but {len(widths)} widths were given; "
                f"the last width was repeated for the rest",
                UserWarning,
                stacklevel=3,
            )
            widths += [widths[-1]] * (n_columns - len(widths))

        # Scale down to the text width when the total overruns it, keeping the
        # proportions the caller asked for. A table narrower than the page is
        # left narrow -- only the all-beams summaries are meant to span it.
        usable = self.usable_width.emu
        total = sum(w.emu for w in widths)
        if total > usable:
            widths = [Emu(int(w.emu * usable / total)) for w in widths]

        # Word honours cell widths only with autofit off and a fixed layout;
        # without both it stretches every table to the full text width and the
        # widths set below are ignored.
        table.autofit = False
        table.allow_autofit = False
        tbl_pr = table._tbl.tblPr
        for existing in tbl_pr.findall(qn("w:tblLayout")):
            tbl_pr.remove(existing)
        tbl_pr.append(parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>'))

        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width

    def add_table(
        self,
        df: pd.DataFrame,
        column_widths: List[Cm],
        font_size: Optional[float] = 9,
    ) -> None:
        """
        Adds a table to the document, populated with data from a DataFrame.

        Parameters
        ----------
        df : pd.DataFrame
            The DataFrame containing the data for the table.
        column_widths : List[Cm]
            List of column widths for the table.
        font_size : float, optional
            Font size (in points) for table text. Default is 10 pt.

        Returns
        -------
        docx.table.Table
            The created table object.
        """
        from docx.shared import Pt

        df = translate_dataframe(round_for_display(df), self.language)

        # --- Create and style table ---
        table = self.doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
        table.style = "Light Shading"
        self.set_col_widths(table, column_widths)

        # --- Header row ---
        for j in range(df.shape[1]):
            table.cell(0, j).text = str(df.columns[j])

        # --- Data rows ---
        for i in range(df.shape[0]):
            for j in range(df.shape[1]):
                value = df.iat[i, j]
                table.cell(i + 1, j).text = str(value)

        # --- Format all text ---
        if font_size is not None:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(font_size)

        # --- First column not bold ---
        for row in table.rows[1:]:
            first_cell = row.cells[0]
            for run in first_cell.paragraphs[0].runs:
                run.font.bold = False

        # --- Add spacer paragraph ---
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)

        return table

    @property
    def usable_width(self) -> Emu:
        """The text column: page width less both margins."""
        section = self.doc.sections[0]
        return Emu(Emu(section.page_width).emu - Emu(section.left_margin).emu - Emu(section.right_margin).emu)

    def content_widths(self, df: pd.DataFrame, min_cm: float = 0.75) -> List[Cm]:
        """Widths proportional to the longest text in each column.

        The wide all-rows tables -- one line per beam or wall -- have too many
        columns to size by hand, and a hand-written list that falls short of
        the column count is not visibly wrong until it is rendered. Measuring
        the frame cannot fall short.

        The result fills the text column exactly, so these tables still span
        the line, which is what distinguishes them from the per-element ones.
        """
        usable_cm = self.usable_width.cm

        lengths = [max([len(str(column))] + [len(str(value)) for value in df[column]]) for column in df.columns]
        # A floor so a one-character column is still readable, then the rest of
        # the width shared out in proportion to what each column has to show.
        spare = usable_cm - min_cm * len(lengths)
        total = sum(lengths)
        return [Cm(min_cm + spare * length / total) for length in lengths]

    def add_table_data(
        self,
        df: pd.DataFrame,
        column_widths=DETAIL_TABLE_WIDTHS,
        font_size: Optional[float] = None,
    ) -> None:
        """Add a data table. ``font_size`` overrides the document default.

        Wide summary tables need a smaller face than the running text to fit
        the page, so the caller can ask for one.
        """
        self.add_table(df, column_widths, font_size=font_size or self.font_size)

    def _shade_verdicts(self, table: Any, df: pd.DataFrame, column: str) -> None:
        """Shade every cell of ``column`` that holds a verdict, green or red.

        Cells that hold neither mark -- a units row, a blank -- are left alone,
        so the colouring says something wherever it appears.
        """
        column_idx = df.columns.get_loc(column)
        for row_offset in range(df.shape[0]):
            verdict = str(df.iat[row_offset, column_idx])
            if verdict not in (PASS_MARK, FAIL_MARK):
                continue
            passed = verdict == PASS_MARK
            cell = table.rows[row_offset + 1].cells[column_idx]
            fill = "C6EFCE" if passed else "FFC7CE"
            font_color = "006100" if passed else "9C0006"
            cell._element.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill}"/>'))
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor.from_string(font_color)

    def add_table_status(
        self,
        df: pd.DataFrame,
        column_widths: List[Cm],
        status_column: str = VERDICT_COLUMN,
        font_size: Optional[float] = None,
    ) -> None:
        """Add a table whose pass/fail column is shaded green or red.

        The same colours :meth:`add_table_dcr` gives the governing row of a
        detailed result, applied per row to one column instead: a reader scans
        the summary for red rather than reading every tick.
        """
        self.add_table(df, column_widths, font_size=font_size or self.font_size)
        self._shade_verdicts(self.doc.tables[-1], df, status_column)

    def add_table_dcr(self, df: pd.DataFrame) -> None:
        """Add a capacity table, with the governing row shaded green or red.

        The whole row, not one cell: this table ends on the combination that
        governs, and that row is the answer the reader came for.
        """
        self.add_table(df, DETAIL_TABLE_WIDTHS, font_size=self.font_size)
        table = self.doc.tables[-1]

        last_row_idx = df.shape[0]
        dcr_column_idx = 2  # Third column holds the DCR

        dcr_value = float(df.iat[-1, dcr_column_idx])
        if dcr_value < 1:
            shading_color = "C6EFCE"  # Green
            font_color = "006100"
        else:
            shading_color = "FFC7CE"  # Red
            font_color = "9C0006"

        for cell in table.rows[last_row_idx].cells:
            cell._element.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shading_color}"/>'))
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor.from_string(font_color)

    def add_table_min_max(self, df: pd.DataFrame, verdict_column: str = VERDICT_COLUMN) -> None:
        """Add a limit-check table, with its verdict column shaded.

        Same colouring as the summary's Status column: a limit check is a
        pass/fail statement too, and reading a column of ticks for the one
        cross is what the colour saves.
        """
        self.add_table(df, LIMIT_TABLE_WIDTHS, font_size=self.font_size)
        if verdict_column in df.columns:
            self._shade_verdicts(self.doc.tables[-1], df, verdict_column)

    def add_figure(self, fig: "plt.Figure", width: float = 16) -> None:
        """
        Inserts a matplotlib Figure into the Word document.

        Parameters
        ----------
        fig : matplotlib.figure.Figure
            The figure to insert.
        width : float
            Width in centimeters (default 12 cm). Height auto-scales.
        """
        buf = BytesIO()
        fig.savefig(buf, format="png", dpi=200, bbox_inches="tight")
        plt.close(fig)
        buf.seek(0)
        self.doc.add_picture(buf, width=Cm(width))
        buf.close()

        # small space after
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)

    def add_image(self, image_path: str, width: Optional[Cm] = None) -> None:
        """
        Insert an external image into the document.

        Parameters
        ----------
        image_path : str
            Absolute or relative path to the image file.
        width : docx.shared.Cm, optional
            Target width. Height scales to keep aspect ratio.
        """
        # paragraph container so image is not glued to previous element
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()

        if width is not None:
            run.add_picture(image_path, width=width)
        else:
            run.add_picture(image_path)

        # small space after
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)

    def save(self, filename: str) -> None:
        """
        Saves the document to a specified file.

        Parameters
        ----------
        filename : str
            The filename (including path) where the document will be saved.

        Returns
        -------
        None
        """
        self.doc.save(filename)

    #  Examples to run in a Jupyter Notebook
    # formatter = Formatter()
    # display(Markdown(formatter.DCR(0.85)))
    # display(Markdown(formatter.DCR_value(0.85)))
    # display(Markdown(formatter.is_lower(0.85,1)))
    # display(Markdown(formatter.is_greater(0.85,1)))
