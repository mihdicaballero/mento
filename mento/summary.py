from typing import List, Dict, Optional
from pandas import DataFrame
import pandas as pd
import copy
from collections import OrderedDict
from docx.shared import Cm

from mento.material import (
    Concrete,
    SteelBar,
    Concrete_EN_1992_2004,
)
from mento.forces import Forces
from mento.beam import RectangularBeam
from mento import mm, cm, kN, MPa, m, inch, ft, kNm
from mento.node import Node
from mento.results import DocumentBuilder
from mento._version import __version__ as MENTO_VERSION


class BeamSummary:
    def __init__(self, concrete: Concrete, steel_bar: SteelBar, beam_list: DataFrame) -> None:
        self.concrete: Concrete = concrete
        self.steel_bar: SteelBar = steel_bar
        self.beam_list: DataFrame = beam_list
        self.units_row: List[str] = []
        self.data: DataFrame = None
        self.nodes: List[Node] = []
        self._beam_summary: List = []
        self.check_and_process_input()
        self.convert_to_nodes()

    def check_and_process_input(self) -> None:
        # Separate the header, units, and data
        self.units_row = self.beam_list.iloc[0].tolist()  # Second row (units)
        data = self.beam_list.iloc[1:].copy()  # Data rows (after removing the units row)

        # Convert NaN in units to "dimensionless"
        self.units_row = ["" if pd.isna(unit) else unit for unit in self.units_row]

        # Validate the units row
        self.validate_units(self.units_row)

        # Convert NaN to 0 in the data rows.
        # astype(float) first (object -> float, NaN preserved) so the
        # subsequent fillna acts on a float frame and does not trigger the
        # pandas "silent downcasting" FutureWarning.
        data.iloc[:, 2:] = data.iloc[:, 2:].astype(float).fillna(0)
        # Convert specific columns to int and others to float
        columns_to_int = ["ns", "n1", "n2", "n3", "n4"]
        for col in columns_to_int:
            if col in data.columns:
                data[col] = data[col].astype(int)

        # Apply units to the corresponding columns, skipping the first column
        for i in range(1, len(self.units_row)):  # Start from the second column (index 1)
            unit_str = self.units_row[i]
            if unit_str != "":
                unit = self.get_unit_variable(unit_str)
                if isinstance(data.iloc[:, i], pd.Series):
                    # data.iloc[:, i] = data.iloc[:, i].apply(lambda x: x * unit)
                    data.iloc[:, i] = data.iloc[:, i].apply(lambda x: x * unit)

        # Store the processed data
        self.data = data
        # print("Processed Data: Ok")

    def validate_units(self, units_row: List) -> None:
        valid_units = {"m", "mm", "cm", "inch", "ft", "kN", "kNm", ""}
        for unit_str in units_row:
            if unit_str and unit_str not in valid_units:
                raise ValueError(f"Invalid unit '{unit_str}' detected. Allowed units: {valid_units}")
        # print("Processed Units: Ok")

    def get_unit_variable(self, unit_str: str) -> Dict:
        # Map strings to actual unit variables (predefined in the script)
        unit_map = {
            "mm": mm,
            "cm": cm,
            "m": m,
            "in": inch,
            "ft": ft,
            "kN": kN,
            "kNm": kNm,
            "MPa": MPa,
        }
        if unit_str in unit_map:
            return unit_map[unit_str]
        else:
            raise ValueError(f"Unit '{unit_str}' is not recognized.")

    def convert_to_nodes(self) -> None:
        self.nodes = []

        for index, row in self.data.iterrows():
            # Extract forces for each row
            M_y = row["My"]
            N_x = row["Nx"]  # Positive for compression
            V_z = row["Vz"]
            comb = row["Comb."]

            # Ensure these are pint.Quantity objects with correct units
            forces = Forces(label=comb, M_y=M_y, N_x=N_x, V_z=V_z)

            # Extract geometric properties of the beam (width and height)
            width = row["b"]
            height = row["h"]
            c_c = row["cc"]

            # Create a rectangular concrete beam using the extracted values
            beam = RectangularBeam(
                label=row["Label"],
                concrete=self.concrete,
                steel_bar=self.steel_bar,
                width=width,
                height=height,
                c_c=c_c,
            )
            # Set transverse rebar (stirrups) for the beam
            n_stirrups = row["ns"]  # Number of stirrups
            d_b = row["dbs"]  # Diameter of rebar (mm)
            s_l = row["sl"]  # Spacing of stirrups (cm)

            if n_stirrups != 0:
                beam.set_transverse_rebar(n_stirrups=n_stirrups, d_b=d_b, s_l=s_l)

            # Set longitudinal rebar at the bottom if n1 is not 0
            n1 = row["n1"]
            d_b1 = row["db1"]  # Diameter in mm
            n2 = row["n2"]
            d_b2 = row["db2"]  # Diameter in mm
            n3 = row["n3"]
            d_b3 = row["db3"]  # Diameter in mm
            n4 = row["n4"]
            d_b4 = row["db4"]  # Diameter in mm
            if n1 != 0:
                if M_y >= 0 * kNm:
                    beam.set_longitudinal_rebar_bot(n1, d_b1, n2, d_b2, n3, d_b3, n4, d_b4)
                else:
                    beam.set_longitudinal_rebar_top(n1, d_b1, n2, d_b2, n3, d_b3, n4, d_b4)
            # Create a Node for each pair of beam and forces
            node = Node(section=beam, forces=forces)

            # Store the section and its corresponding forces
            self.nodes.append(node)

    def check(self, capacity_check: bool = False) -> DataFrame:
        """
        Perform a check on all beams in the summary.

        Parameters
        ----------
        capacity_check : bool, optional
            If True, resets all forces in the node to zero to perform a capacity check.
            Otherwise, uses the forces currently assigned to the node.

        Returns
        -------
        DataFrame
            A DataFrame with the results of the check.
        """

        results_list = []
        for node in self.nodes:
            beam: RectangularBeam = node.section  # type: ignore
            original_forces = [copy.deepcopy(force) for force in node.get_forces_list()]

            rebar_v = (
                "-"
                if beam._stirrup_n == 0
                else f"{int(beam._stirrup_n)}eØ{int(beam._stirrup_d_b.to('mm').magnitude)}/{int(beam._stirrup_s_l.to('cm').magnitude)}"
            )  # noqa: E501
            rebar_f_top = (
                "-"
                if beam._n1_t == 0
                else (
                    f"{beam._format_longitudinal_rebar_string(beam._n1_t, beam._d_b1_t, beam._n2_t, beam._d_b2_t)}"
                    + (
                        f" ++ {beam._format_longitudinal_rebar_string(beam._n3_t, beam._d_b3_t, beam._n4_t, beam._d_b4_t)}"
                        if beam._n3_t != 0
                        else ""
                    )
                )
            )
            rebar_f_bot = (
                "-"
                if beam._n1_b == 0
                else (
                    f"{beam._format_longitudinal_rebar_string(beam._n1_b, beam._d_b1_b, beam._n2_b, beam._d_b2_b)}"
                    + (
                        f" ++ {beam._format_longitudinal_rebar_string(beam._n3_b, beam._d_b3_b, beam._n4_b, beam._d_b4_b)}"
                        if beam._n3_b != 0
                        else ""
                    )
                )
            )

            if capacity_check:
                # Remove all forces assignments
                node.clear_forces()
                # Create empty force
                empty_force = Forces()
                node.add_forces(empty_force)
                # Perform the shear check
                shear_results = node.check_shear()
                node.check_flexure()
                # Common data
                common_data = {
                    "Beam": beam.label,
                    "b": beam.width.magnitude,
                    "h": beam.height.magnitude,
                    "As,top": rebar_f_top,
                    "As,bot": rebar_f_bot,
                    "Av": rebar_v,
                    "As,top,real": round(beam._A_s_top.to("cm**2").magnitude, 1),
                    "As,bot,real": round(beam._A_s_bot.to("cm**2").magnitude, 1),
                    "Av,real": round(shear_results["Av"][1], 1),
                }

                common_units = {
                    "Beam": "",
                    "b": "cm",
                    "h": "cm",
                    "As,top": "",
                    "As,bot": "",
                    "Av": "",
                    "As,top,real": "cm²",
                    "As,bot,real": "cm²",
                    "Av,real": "cm²/m",
                }

                # Code-specific data #TODO
                if isinstance(self.concrete, Concrete_EN_1992_2004):
                    code_specific_data = {
                        "MRd,top": round(beam._M_Rd_top.to("kN*m").magnitude, 1),
                        "MRd,bot": round(beam._M_Rd_bot.to("kN*m").magnitude, 1),
                        "VRd": shear_results["VRd"][1],
                    }
                    code_specific_units = {
                        "MRd,top": "kNm",
                        "MRd,bot": "kNm",
                        "VRd": "kN",
                    }
                else:  # ACI 318-19 or CIRSOC 201-25 design code
                    code_specific_data = {
                        "ØMn,top": round(beam._phi_M_n_top.to("kN*m").magnitude, 1),
                        "ØMn,bot": round(beam._phi_M_n_bot.to("kN*m").magnitude, 1),
                        "ØVn": shear_results["ØVn"][1],
                    }
                    code_specific_units = {
                        "ØMn,top": "kNm",
                        "ØMn,bot": "kNm",
                        "ØVn": "kN",
                    }

                # Merge data dictionaries
                merged_data = {**common_data, **code_specific_data}
                results_dict = merged_data

                # Merge units dictionaries
                merged_units = {**common_units, **code_specific_units}
                units_row = pd.DataFrame([OrderedDict({**merged_units})])

                # Restore the original forces after capacity check
                node.clear_forces()
                node.add_forces(original_forces)
            else:
                # Perform the shear check
                shear_results = node.check_shear().iloc[1:].reset_index(drop=True)  # Skip the first row (units)
                flexure_results = node.check_flexure().iloc[1:].reset_index(drop=True)  # Skip the first row (units)
                # Common data that doesn't change between codes
                common_data = {
                    "Beam": beam.label,
                    "b": int(beam.width.magnitude),
                    "h": int(beam.height.magnitude),
                    "cc": int(beam.c_c.magnitude),
                    "As,top": rebar_f_top,
                    "As,bot": rebar_f_bot,
                    "Av": rebar_v,
                    "As,req,top": round(beam._A_s_req_top.to("cm**2").magnitude, 1),
                    "As,req,bot": round(beam._A_s_req_bot.to("cm**2").magnitude, 1),
                    "Av,req": round(shear_results["Av,req"][0], 2),
                    "Av,real": round(shear_results["Av"][0], 2),
                    "DCRb,top": round(beam._DCRb_top, 3),
                    "DCRb,bot": round(beam._DCRb_bot, 3),
                    "DCRv": shear_results["DCR"][0],
                }

                common_units = {
                    "Beam": "",
                    "b": "cm",
                    "h": "cm",
                    "cc": "mm",
                    "As,top": "",
                    "As,bot": "",
                    "Av": "",
                    "As,req,top": "cm²/m",
                    "As,req,bot": "cm²/m",
                    "Av,req": "cm²/m",
                    "Av,real": "cm²/m",
                    "DCRb,top": "",
                    "DCRb,bot": "",
                    "DCRv": "",
                }

                # Code-specific data
                if isinstance(self.concrete, Concrete_EN_1992_2004):  # TODO
                    code_specific_data = {
                        "MEd": round(flexure_results["MEd"][0], 1),
                        "VEd": round(shear_results["VEd,2"][0], 1),
                        "NEd": round(shear_results["NEd"][0], 1),
                        "MRd,top": round(beam._M_Rd_top.to("kN*m").magnitude, 1),
                        "MRd,bot": round(beam._M_Rd_bot.to("kN*m").magnitude, 1),
                        "VRd": round(shear_results["VRd"][0], 1),
                    }
                    code_specific_units = {
                        "MEd": "kNm",
                        "VEd": "kN",
                        "NEd": "kN",
                        "MRd,top": "kNm",
                        "MRd,bot": "kNm",
                        "VRd": "kN",
                    }
                else:  # ACI 318-19 or CIRSOC 201-25 design code
                    code_specific_data = {
                        "Mu": round(flexure_results["Mu"][0], 1),
                        "Vu": round(shear_results["Vu"][0], 1),
                        "Nu": round(shear_results["Nu"][0], 1),
                        "ØMn,top": round(beam._phi_M_n_top.to("kN*m").magnitude, 1),
                        "ØMn,bot": round(beam._phi_M_n_bot.to("kN*m").magnitude, 1),
                        "ØVn": round(shear_results["ØVn"][0], 1),
                    }
                    code_specific_units = {
                        "Mu": "kNm",
                        "Vu": "kN",
                        "Nu": "kN",
                        "ØMn,top": "kNm",
                        "ØMn,bot": "kNm",
                        "ØVn": "kN",
                    }

                # Assemble results_dict and units_row from the split dicts
                results_dict = OrderedDict({**common_data, **code_specific_data})
                units_row = pd.DataFrame([OrderedDict({**common_units, **code_specific_units, "Status": ""})])

                # Determine status
                dcr_values = [results_dict["DCRb,top"], results_dict["DCRb,bot"], results_dict["DCRv"]]
                all_dcrs_ok = all(v < 1 for v in dcr_values)
                results_dict["Status"] = "✅" if all_dcrs_ok else "❌"

            # Add the results to the list
            results_list.append(results_dict)

        # Convert results list into a DataFrame
        results_df = pd.DataFrame(results_list)

        # Combine the units row with the results DataFrame
        final_df = pd.concat([units_row, results_df], ignore_index=True)
        return final_df

    def design(self) -> DataFrame:
        """
        Run design for all beams in the summary.
        Fills in the rebar columns (n1–n4, db1–db4, ns, dbs, sl)
        with the suggested designs for shear and flexure.

        Returns
        -------
        DataFrame
            A copy of the beam summary with designed reinforcement filled in.
        """

        # Copy the processed data to avoid overwriting self.data
        design_df = self.data.copy()
        design_df = self.data.reset_index(drop=True).copy()

        for i, node in enumerate(self.nodes):
            beam: RectangularBeam = node.section  # type: ignore
            forces = node.forces[0]

            # --- FLEXURE DESIGN ---
            node.design_flexure()
            if forces.M_y.magnitude >= 0:  # tension at bottom face
                flex_result = beam.flexure_design_results_bot
                # print(beam.flexure_design_results_bot)
                design_df.loc[i, "n1"] = flex_result["n_1"]
                design_df.loc[i, "db1"] = flex_result["d_b1"]
                design_df.loc[i, "n2"] = flex_result["n_2"]
                design_df.loc[i, "db2"] = flex_result["d_b2"] if flex_result["d_b2"] is not None else 0
                design_df.loc[i, "n3"] = flex_result["n_3"]
                design_df.loc[i, "db3"] = flex_result["d_b3"] if flex_result["d_b3"] is not None else 0
                design_df.loc[i, "n4"] = flex_result["n_4"]
                design_df.loc[i, "db4"] = flex_result["d_b4"] if flex_result["d_b4"] is not None else 0
            else:  # tension at top face
                flex_result = beam.flexure_design_results_top
                design_df.loc[i, "n1"] = flex_result["n_1"]
                design_df.loc[i, "db1"] = flex_result["d_b1"]
                design_df.loc[i, "n2"] = flex_result["n_2"]
                design_df.loc[i, "db2"] = flex_result["d_b2"] if flex_result["d_b2"] is not None else 0
                design_df.loc[i, "n3"] = flex_result["n_3"]
                design_df.loc[i, "db3"] = flex_result["d_b3"] if flex_result["d_b3"] is not None else 0
                design_df.loc[i, "n4"] = flex_result["n_4"]
                design_df.loc[i, "db4"] = flex_result["d_b4"] if flex_result["d_b4"] is not None else 0

            # --- SHEAR DESIGN ---
            node.design_shear()
            shear_row = beam.shear_design_results.iloc[0]  # take best row
            design_df.loc[i, "ns"] = int(shear_row["n_stir"])
            design_df.loc[i, "dbs"] = shear_row["d_b"]
            design_df.loc[i, "sl"] = shear_row["s_l"]

        # store for export
        self.design_data = design_df

        print("✅ Beam design completed for all beams in Summary.")
        return design_df

    def shear_results(self, index: Optional[int] = None, capacity_check: bool = False) -> DataFrame:
        """
        Get shear results for one or all beams.
        Includes a units row only once at the top.
        """
        if index is not None:
            if index - 1 >= len(self.nodes):
                raise IndexError(f"Index {index} is out of range for the beam list.")
            node = self.nodes[max(index - 1, 0)]
            df = self._process_beam_for_check(node, "shear", capacity_check)
            units_row = node.section._get_units_row_shear()  # type: ignore
            df_all = pd.concat([units_row, df], ignore_index=True)
        else:
            # For all nodes, only include the units row once
            results = []
            units_row_added = False
            for item in self.nodes:
                df = self._process_beam_for_check(item, "shear", capacity_check)
                if not units_row_added:
                    units_row = item.section._get_units_row_shear()  # type: ignore
                    results.append(units_row)
                    units_row_added = True
                results.append(df)

            df_all = pd.concat(results, ignore_index=True)

        # Separate the units row (first row) from the data
        units_row = df_all.iloc[[0]]  # DataFrame with 1 row
        data_rows = df_all.iloc[1:].copy()

        # Recombine units + data
        df_final = pd.concat([units_row, data_rows], ignore_index=True)
        return df_final

    def flexure_results(self, index: Optional[int] = None, capacity_check: bool = False) -> DataFrame:
        """
        Get flexure results for one or all beams.
        Includes a units row only once at the top.
        """
        if index is not None:
            if index - 1 >= len(self.nodes):
                raise IndexError(f"Index {index} is out of range for the beam list.")
            node = self.nodes[max(index - 1, 0)]
            df = self._process_beam_for_check(node, "flexure", capacity_check)
            units_row = node.section._get_units_row_flexure()  # type: ignore
            df_all = pd.concat([units_row, df], ignore_index=True)
        else:
            # For all nodes, only include the units row once
            results = []
            units_row_added = False
            for item in self.nodes:
                df = self._process_beam_for_check(item, "flexure", capacity_check)
                if not units_row_added:
                    units_row = item.section._get_units_row_flexure()  # type: ignore
                    results.append(units_row)
                    units_row_added = True
                results.append(df)

            df_all = pd.concat(results, ignore_index=True)

        # Separate the units row (first row) from the data
        units_row = df_all.iloc[[0]]  # DataFrame with 1 row
        data_rows = df_all.iloc[1:].copy()

        # Recombine units + data
        df_final = pd.concat([units_row, data_rows], ignore_index=True)
        return df_final

    def _process_beam_for_check(self, node: Node, check_type: str, capacity_check: bool) -> DataFrame:
        """
        Shared method to process beam for either shear or flexure checks.

        :param node: Node object to check
        :param check_type: Either 'shear' or 'flexure'
        :param capacity_check: If True, performs capacity check (resets forces)
        :return: Results DataFrame
        """
        original_forces = [copy.deepcopy(force) for force in node.get_forces_list()]

        if capacity_check:
            node.clear_forces()
            node.add_forces(Forces())  # Add empty force

        # Run the appropriate check
        if check_type == "shear":
            results = node.check_shear().iloc[1:].reset_index(drop=True)
        elif check_type == "flexure":
            results = node.check_flexure().iloc[1:].reset_index(drop=True)
        else:
            raise ValueError("check_type must be either 'shear' or 'flexure'")

        # Add code-specific capacity columns after a capacity flexure check
        if capacity_check and check_type == "flexure":
            beam: RectangularBeam = node.section  # type: ignore
            if isinstance(self.concrete, Concrete_EN_1992_2004):
                results["MRd,top"] = round(beam._M_Rd_top.to("kN*m").magnitude, 1)
                results["MRd,bot"] = round(beam._M_Rd_bot.to("kN*m").magnitude, 1)
            else:
                results["ØMn,top"] = round(beam._phi_M_n_top.to("kN*m").magnitude, 1)
                results["ØMn,bot"] = round(beam._phi_M_n_bot.to("kN*m").magnitude, 1)

        # Restore original forces if we did a capacity check
        if capacity_check:
            node.clear_forces()
            node.add_forces(original_forces)
        return results

    # ------------------------------------------------------------
    # EXCEL I/O HELPERS FOR DESIGN / CHECK WORKFLOW
    # ------------------------------------------------------------

    def export_design(self, path: str) -> None:
        """
        Export current beam list (including units row) to Excel.
        Typically used after design() to create an editable file.

        Parameters
        ----------
        path : str
            Path to the Excel file to create.
        """
        if not hasattr(self, "design_data"):
            raise AttributeError("No design data found. Run .design() before exporting.")

        df_numeric = self.design_data.copy()
        for col in df_numeric.columns:
            df_numeric[col] = df_numeric[col].apply(lambda x: x.magnitude if hasattr(x, "magnitude") else x)
        # Recombine units + data before exporting
        df_export = pd.concat(
            [
                pd.DataFrame([self.units_row], columns=self.beam_list.columns),
                df_numeric,
            ],
            ignore_index=True,
        )
        df_export.to_excel(path, index=False)
        print(f"✅ Beam design exported to {path}")

    def import_design(self, path: str) -> None:
        """
        Import an Excel file containing edited or designed rebar information.
        Updates self.beam_list, reprocesses units, and rebuilds nodes.

        Parameters
        ----------
        path : str
            Path to the Excel file to import.
        """
        beam_df = pd.read_excel(path)

        # Replace the original beam_list and re-process input
        self.beam_list = beam_df
        self.check_and_process_input()
        self.convert_to_nodes()
        print("✅ Beam design imported and summary data updated.")

    def results_detailed_doc(self, index: int = 1) -> None:
        """
        Export detailed results to Word document.
        Shows detailed shear/flexure for one beam, then summary tables for all.

        Parameters
        ----------
        index : int
            1-based index of the beam to show detailed results for (default: 1)
        """

        if index < 1 or index > len(self.nodes):
            raise IndexError(f"Index {index} out of range. Valid: 1 to {len(self.nodes)}")

        node = self.nodes[index - 1]
        beam: RectangularBeam = node.section  # type: ignore

        # Run checks if not already done
        node.check_flexure()
        node.check_shear()

        # Create document with smaller font
        doc_builder = DocumentBuilder(title="Beam Summary Analysis", font_size=8)
        doc_builder.add_heading("Beam Summary Analysis", level=1)
        doc_builder.add_text(f"Made with mento {MENTO_VERSION}. Design code: {self.concrete.design_code}")
        doc_builder.add_text(
            "This report presents the detailed results for the first beam of the summary, followed by summary tables for all beams."
        )

        # --- DETAILED FLEXURE RESULTS FOR SELECTED BEAM ---
        doc_builder.add_heading(f"Beam {beam.label} flexure check", level=2)

        # Build dataframes same as flexure_results_detailed_doc
        top_result_data = beam._limiting_case_flexure_top_details["flexure_capacity_top"]
        bot_result_data = beam._limiting_case_flexure_bot_details["flexure_capacity_bot"]
        forces_result = {
            "Design forces": ["Top max moment", "Bottom max moment"],
            "Variable": ["Mu,top", "Mu,bot"],
            "Value": [
                round(beam._limiting_case_flexure_top_details["forces"]["Value"][0], 2),
                round(beam._limiting_case_flexure_bot_details["forces"]["Value"][1], 2),
            ],
            "Unit": ["kNm", "kNm"],
        }
        min_max_result = {
            "Check": [
                "Min/Max As rebar top",
                "Minimum spacing top",
                "Min/Max As rebar bottom",
                "Minimum spacing bottom",
            ],
            "Unit": ["cm²", "mm", "cm²", "mm"],
            "Value": [
                round(beam._limiting_case_flexure_top_details["min_max"]["Value"][0], 2),
                round(beam._limiting_case_flexure_top_details["min_max"]["Value"][1], 2),
                round(beam._limiting_case_flexure_bot_details["min_max"]["Value"][2], 2),
                round(beam._limiting_case_flexure_bot_details["min_max"]["Value"][3], 2),
            ],
            "Min.": [
                round(beam._limiting_case_flexure_top_details["min_max"]["Min."][0], 2),
                beam._limiting_case_flexure_top_details["min_max"]["Min."][1],
                round(beam._limiting_case_flexure_bot_details["min_max"]["Min."][2], 2),
                beam._limiting_case_flexure_bot_details["min_max"]["Min."][3],
            ],
            "Max.": [
                round(beam._limiting_case_flexure_top_details["min_max"]["Max."][0], 2),
                "",
                round(beam._limiting_case_flexure_bot_details["min_max"]["Max."][2], 2),
                "",
            ],
            "Ok?": [
                beam._limiting_case_flexure_top_details["min_max"]["Ok?"][0],
                beam._limiting_case_flexure_top_details["min_max"]["Ok?"][1],
                beam._limiting_case_flexure_bot_details["min_max"]["Ok?"][2],
                beam._limiting_case_flexure_bot_details["min_max"]["Ok?"][3],
            ],
        }

        df_flex_materials = pd.DataFrame(beam._materials_flexure)
        df_flex_geometry = pd.DataFrame(beam._geometry_flexure)
        df_flex_forces = pd.DataFrame(forces_result)
        df_flex_min_max = pd.DataFrame(min_max_result)
        df_flex_capacity_top = pd.DataFrame(top_result_data)
        df_flex_capacity_bot = pd.DataFrame(bot_result_data)

        doc_builder.add_heading("Materials", level=3)
        doc_builder.add_table_data(df_flex_materials)
        doc_builder.add_table_data(df_flex_geometry)
        doc_builder.add_table_data(df_flex_forces)
        doc_builder.add_heading("Limit Checks", level=3)
        doc_builder.add_table_data(df_flex_min_max)
        doc_builder.add_heading("Flexural Capacity Top", level=3)
        doc_builder.add_table_dcr(df_flex_capacity_top)
        doc_builder.add_heading("Flexural Capacity Bottom", level=3)
        doc_builder.add_table_dcr(df_flex_capacity_bot)

        # --- DETAILED SHEAR RESULTS FOR SELECTED BEAM ---
        doc_builder.add_heading(f"Beam {beam.label} shear check", level=2)

        result_data = beam._limiting_case_shear_details
        df_shear_materials = pd.DataFrame(beam._materials_shear)
        df_shear_geometry = pd.DataFrame(beam._geometry_shear)
        df_shear_forces = pd.DataFrame(result_data["forces"])
        df_shear_reinforcement = pd.DataFrame(result_data["shear_reinforcement"])
        df_shear_min_max = pd.DataFrame(result_data["min_max"])
        df_shear_concrete = pd.DataFrame(result_data["shear_concrete"])

        doc_builder.add_heading("Materials", level=3)
        doc_builder.add_table_data(df_shear_materials)
        doc_builder.add_table_data(df_shear_geometry)
        doc_builder.add_table_data(df_shear_forces)
        doc_builder.add_heading("Limit checks", level=3)
        doc_builder.add_table_min_max(df_shear_min_max)
        doc_builder.add_heading("Design checks", level=3)
        doc_builder.add_table_data(df_shear_reinforcement)
        doc_builder.add_table_dcr(df_shear_concrete)

        # --- SUMMARY TABLES FOR ALL BEAMS ---
        doc_builder.add_heading("Summary - All Beams", level=2)
        doc_builder.add_heading("Beam Data", level=3)
        beam_data_out = self.beam_list.fillna("")
        doc_builder.add_table_data(
            beam_data_out,
            column_widths=[
                Cm(1),
                Cm(1),
                Cm(1),
                Cm(1),
                Cm(1),
                Cm(1.5),
                Cm(1.5),
                Cm(1),
                Cm(1),
                Cm(1),
                Cm(1),
                Cm(1),
                Cm(1),
                Cm(1),
                Cm(1),
            ],
        )

        doc_builder.add_heading("Flexure Results", level=3)
        df_flex_all = self.flexure_results(capacity_check=False)
        doc_builder.add_table_data(
            df_flex_all,
            column_widths=[Cm(2), Cm(2), Cm(2), Cm(2), Cm(2), Cm(1.5), Cm(1.5), Cm(1.5), Cm(1.5), Cm(1.5), Cm(1.5)],
        )

        doc_builder.add_heading("Shear Results", level=3)
        df_shear_all = self.shear_results(capacity_check=False)

        if isinstance(self.concrete, Concrete_EN_1992_2004):
            # Remove Eurocode-specific columns
            cols_to_remove = ["VEd,1≤VRd,max", "VEd,2≤VRd"]
            df_shear_all = df_shear_all.drop(columns=[c for c in cols_to_remove if c in df_shear_all.columns])

            doc_builder.add_table_data(
                df_shear_all,
                column_widths=[
                    Cm(2),
                    Cm(2),
                    Cm(2),
                    Cm(2),
                    Cm(2),
                    Cm(1.5),
                    Cm(1.5),
                    Cm(1.5),
                    Cm(1.5),
                    Cm(1.5),
                    Cm(1.5),
                    Cm(2),
                ],
            )
        else:
            doc_builder.add_table_data(
                df_shear_all,
                column_widths=[
                    Cm(2),
                    Cm(2),
                    Cm(2),
                    Cm(2),
                    Cm(2),
                    Cm(1.5),
                    Cm(1.5),
                    Cm(1.5),
                    Cm(1.5),
                    Cm(1.5),
                    Cm(1.5),
                    Cm(2),
                    Cm(2),
                    Cm(2),
                ],
            )

        doc_builder.add_heading("Design Check Summary", level=3)
        df_check = self.check()
        doc_builder.add_table_data(
            df_check, column_widths=[Cm(1.7), Cm(1), Cm(1), Cm(3), Cm(3), Cm(2), Cm(1.5), Cm(1.5), Cm(1.2), Cm(1.2)]
        )

        # Save
        doc_builder.save(f"Beam_Summary_{self.concrete.design_code}.docx")
        print(f"✅ Results exported to Beam_Summary_{self.concrete.design_code}.docx")
