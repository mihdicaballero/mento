"""Word reports for a summary of many elements.

The document assembly used to live on ``BeamSummary`` and ``ShearWallSummary``.
Phase 3 of the architecture roadmap moves it here; both classes keep a one-line
``results_detailed_doc`` delegation, so nothing calling them changes.

These are module functions taking the summary, the same shape the design-code
modules use. Sharing one namespace, they are named after the element family
rather than after the method.
"""

from __future__ import annotations

from typing import TYPE_CHECKING, Any, Dict, Optional, cast

import pandas as pd
from docx.shared import Cm

from mento._version import __version__ as MENTO_VERSION
from mento.codes.registry import design_code
from mento.results import DocumentBuilder

if TYPE_CHECKING:
    from mento.beam import RectangularBeam
    from mento.beam_summary import BeamSummary
    from mento.shear_wall import ShearWall
    from mento.shear_wall_summary import ShearWallSummary


#: The all-beams tables are much wider than the running text, so they are set
#: a point smaller to keep every column on the page.
SUMMARY_FONT_SIZE = 7

#: What the Beam Data table lists: the section and the bars it carries. The
#: input frame also holds the position and the demands, which belong to the
#: per-combination tables rather than to a list of sections.
BEAM_DATA_COLUMNS = (
    "Label",
    "b",
    "h",
    "cc",
    "ns",
    "dbs",
    "sl",
    "n1",
    "db1",
    "n2",
    "db2",
    "n3",
    "db3",
    "n4",
    "db4",
)

#: The label needs room for a beam name and the dimensions for two digits; the
#: eleven rebar columns hold a count or a diameter and no more.
BEAM_DATA_WIDTHS = [Cm(2), Cm(1), Cm(1), Cm(1)] + [Cm(0.9)] * 11

#: Widths for the two per-combination summaries, one entry per column, set
#: against the rendered document rather than computed. Both design codes leave
#: the same three capacity ticks out of these tables, so both end up the same
#: shape -- ten columns and twelve -- and one list each serves them. If a code
#: ever drops a different set, `add_table` warns that its list fell short
#: rather than quietly repeating the last width.
FLEXURE_SUMMARY_WIDTHS = [
    Cm(2),
    Cm(4),
    Cm(1.3),
    Cm(1.2),
    Cm(1.6),
    Cm(1.6),
    Cm(1.2),
    Cm(1.2),
    Cm(1.2),
    Cm(1),
]

SHEAR_SUMMARY_WIDTHS = [
    Cm(2),
    Cm(4),
    Cm(1.2),
    Cm(1.2),
    Cm(1.1),
    Cm(1.2),
    Cm(1.2),
    Cm(1.2),
    Cm(1.2),
    Cm(1.2),
    Cm(1.4),
    Cm(1),
]

#: The closing table: the section, its bars, the governing demands, the three
#: DCRs and the verdict.
CHECK_SUMMARY_WIDTHS = [
    Cm(2),
    Cm(1),
    Cm(1),
    Cm(1.4),
    Cm(1.4),
    Cm(1.4),
    Cm(1.2),
    Cm(1.2),
    Cm(1.2),
    Cm(1.5),
    Cm(1.5),
    Cm(1.1),
    Cm(1.0),
]


def _without_dropped_columns(self: "BeamSummary", df: pd.DataFrame) -> pd.DataFrame:
    """Drop the columns the active design code keeps out of its Word summary."""
    dropped = design_code(self.concrete).summary_drop_columns
    return df.drop(columns=[column for column in dropped if column in df.columns])


def _details(value: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    """The detail tables a check leaves on the element.

    Typed ``Optional`` on the element because one that was never checked has
    none; both functions here run the checks before reading them.
    """
    return cast(Dict[str, Any], value)


def beam_summary_doc(self: "BeamSummary", index: int = 1) -> None:
    """
    Export detailed results to Word document.
    Shows detailed shear/flexure for one beam, then summary tables for all.

    Parameters
    ----------
    index : int
        1-based index of the beam to show detailed results for (default: 1)
    """

    if index < 1 or index > len(self.nodes):
        raise IndexError(f"Index {index} out of range. Valid: 1 to {len(self.nodes)}")

    node = self.nodes[index - 1]
    beam: RectangularBeam = node.section  # type: ignore

    # Run checks if not already done
    node.check_flexure()
    node.check_shear()

    # Create document with smaller font
    doc_builder = DocumentBuilder(title="Beam Summary Analysis", font_size=8)
    doc_builder.add_heading("Beam Summary Analysis", level=1)
    doc_builder.add_text(f"Made with mento {MENTO_VERSION}. Design code: {self.concrete.design_code}")
    doc_builder.add_text(
        "This report presents the detailed results for the first beam of the summary, followed by summary tables for all beams."
    )

    # --- DETAILED FLEXURE RESULTS FOR SELECTED BEAM ---
    doc_builder.add_heading(f"Beam {beam.label} flexure check", level=2)

    # Build dataframes same as flexure_results_detailed_doc
    top_details = _details(beam._limiting_case_flexure_top_details)
    bot_details = _details(beam._limiting_case_flexure_bot_details)
    top_result_data = top_details["flexure_capacity_top"]
    bot_result_data = bot_details["flexure_capacity_bot"]
    forces_result = {
        "Design forces": ["Top max moment", "Bottom max moment"],
        "Variable": [
            beam._flexure_symbols["demand_top"],
            beam._flexure_symbols["demand_bot"],
        ],
        "Value": [
            round(top_details["forces"]["Value"][0], 2),
            round(bot_details["forces"]["Value"][1], 2),
        ],
        "Unit": ["kNm", "kNm"],
    }
    min_max_result = {
        "Check": [
            "Min/Max As rebar top",
            "Minimum spacing top",
            "Min/Max As rebar bottom",
            "Minimum spacing bottom",
        ],
        "Unit": ["cm²", "mm", "cm²", "mm"],
        "Value": [
            round(top_details["min_max"]["Value"][0], 2),
            round(top_details["min_max"]["Value"][1], 2),
            round(bot_details["min_max"]["Value"][2], 2),
            round(bot_details["min_max"]["Value"][3], 2),
        ],
        "Min.": [
            round(top_details["min_max"]["Min."][0], 2),
            top_details["min_max"]["Min."][1],
            round(bot_details["min_max"]["Min."][2], 2),
            bot_details["min_max"]["Min."][3],
        ],
        "Max.": [
            round(top_details["min_max"]["Max."][0], 2),
            "",
            round(bot_details["min_max"]["Max."][2], 2),
            "",
        ],
        "Ok?": [
            top_details["min_max"]["Ok?"][0],
            top_details["min_max"]["Ok?"][1],
            bot_details["min_max"]["Ok?"][2],
            bot_details["min_max"]["Ok?"][3],
        ],
    }

    df_flex_materials = pd.DataFrame(beam._materials_flexure)
    df_flex_geometry = pd.DataFrame(beam._geometry_flexure)
    df_flex_forces = pd.DataFrame(forces_result)
    df_flex_min_max = pd.DataFrame(min_max_result)
    df_flex_capacity_top = pd.DataFrame(top_result_data)
    df_flex_capacity_bot = pd.DataFrame(bot_result_data)

    doc_builder.add_heading("Materials", level=3)
    doc_builder.add_table_data(df_flex_materials)
    doc_builder.add_table_data(df_flex_geometry)
    doc_builder.add_table_data(df_flex_forces)
    doc_builder.add_heading("Limit checks", level=3)
    doc_builder.add_table_min_max(df_flex_min_max)
    doc_builder.add_heading("Flexural Capacity Top", level=3)
    doc_builder.add_table_dcr(df_flex_capacity_top)
    doc_builder.add_heading("Flexural Capacity Bottom", level=3)
    doc_builder.add_table_dcr(df_flex_capacity_bot)

    # --- DETAILED SHEAR RESULTS FOR SELECTED BEAM ---
    doc_builder.add_heading(f"Beam {beam.label} shear check", level=2)

    result_data = _details(beam._limiting_case_shear_details)
    df_shear_materials = pd.DataFrame(beam._materials_shear)
    df_shear_geometry = pd.DataFrame(beam._geometry_shear)
    df_shear_forces = pd.DataFrame(result_data["forces"])
    df_shear_reinforcement = pd.DataFrame(result_data["shear_reinforcement"])
    df_shear_min_max = pd.DataFrame(result_data["min_max"])
    df_shear_concrete = pd.DataFrame(result_data["shear_concrete"])

    doc_builder.add_heading("Materials", level=3)
    doc_builder.add_table_data(df_shear_materials)
    doc_builder.add_table_data(df_shear_geometry)
    doc_builder.add_table_data(df_shear_forces)
    doc_builder.add_heading("Limit checks", level=3)
    doc_builder.add_table_min_max(df_shear_min_max)
    doc_builder.add_heading("Design checks", level=3)
    doc_builder.add_table_data(df_shear_reinforcement)
    doc_builder.add_table_dcr(df_shear_concrete)

    # --- SUMMARY TABLES FOR ALL BEAMS ---
    doc_builder.add_heading("Summary - All Beams", level=2)
    doc_builder.add_heading("Beam Data", level=3)
    # Geometry and reinforcement only: the demands each beam was checked for
    # are reported by the flexure and shear tables below, per combination,
    # which is where they mean something.
    beam_data_out = self.beam_list.fillna("")[list(BEAM_DATA_COLUMNS)]
    doc_builder.add_table_data(
        beam_data_out,
        column_widths=BEAM_DATA_WIDTHS,
        font_size=SUMMARY_FONT_SIZE,
    )

    # The label and the combination name are the only columns holding words;
    # the rest hold a number each and share what is left. Shared rather than
    # listed because the column count is the code's: each code drops a
    # different set of columns below, so one hand-written list cannot serve
    # both without going quietly out of step the next time one is dropped.
    doc_builder.add_heading("Flexure Results", level=3)
    df_flex_all = _without_dropped_columns(self, self.flexure_results(capacity_check=False))
    doc_builder.add_table_data(
        df_flex_all,
        column_widths=FLEXURE_SUMMARY_WIDTHS,
        font_size=SUMMARY_FONT_SIZE,
    )

    doc_builder.add_heading("Shear Results", level=3)
    df_shear_all = _without_dropped_columns(self, self.shear_results(capacity_check=False))
    doc_builder.add_table_data(
        df_shear_all,
        column_widths=SHEAR_SUMMARY_WIDTHS,
        font_size=SUMMARY_FONT_SIZE,
    )

    doc_builder.add_heading("Design Check Summary", level=3)
    # Printed as `check()` returns it. The table used to select a subset here,
    # which meant the notebook and the report disagreed about what the summary
    # is; `check()` carries the shorter set now and this prints it.
    df_check = self.check()
    doc_builder.add_table_status(
        df_check,
        column_widths=CHECK_SUMMARY_WIDTHS,
        font_size=SUMMARY_FONT_SIZE,
    )

    # Save
    doc_builder.save(f"Beam_Summary_{self.concrete.design_code}.docx")
    print(f"✅ Results exported to Beam_Summary_{self.concrete.design_code}.docx")


def wall_summary_doc(self: "ShearWallSummary", index: int = 1) -> None:
    if index < 1 or index > len(self.nodes):
        raise IndexError(f"Index {index} out of range. Valid: 1 to {len(self.nodes)}")

    node = self.nodes[index - 1]
    wall: ShearWall = node.section  # type: ignore

    node.check_shear()

    doc_builder = DocumentBuilder(title="Shear Wall Summary Analysis", font_size=8)
    doc_builder.add_heading("Shear Wall Summary Analysis", level=1)
    doc_builder.add_text(f"Made with mento {MENTO_VERSION}. Design code: {self.concrete.design_code}")

    # --- DETAILED RESULTS FOR SELECTED WALL ---
    doc_builder.add_heading(f"Wall {wall.level} - {wall.label} shear check", level=2)

    result_data = _details(wall._limiting_case_shear_details)
    df_materials = pd.DataFrame(wall._materials_shear_wall)
    df_geometry = pd.DataFrame(wall._geometry_shear_wall)
    df_forces = pd.DataFrame(result_data["forces"])
    df_min_max = pd.DataFrame(result_data["min_max"])
    df_capacity = pd.DataFrame(result_data["shear_capacity"])

    doc_builder.add_heading("Materials", level=3)
    doc_builder.add_table_data(df_materials)
    doc_builder.add_table_data(df_geometry)
    doc_builder.add_table_data(df_forces)
    doc_builder.add_heading("Limit checks", level=3)
    doc_builder.add_table_min_max(df_min_max)
    doc_builder.add_heading("Design checks", level=3)
    doc_builder.add_table_dcr(df_capacity)

    # --- SUMMARY TABLES FOR ALL WALLS ---
    doc_builder.add_heading("Summary - All Walls", level=2)

    doc_builder.add_heading("Wall Data", level=3)
    wall_data_out = self.wall_list.fillna("")
    doc_builder.add_table_data(
        wall_data_out, column_widths=doc_builder.content_widths(wall_data_out), font_size=SUMMARY_FONT_SIZE
    )

    doc_builder.add_heading("Shear Results", level=3)
    df_shear_all = self.shear_results()
    cols_to_drop = [c for c in ["Vu≤ØVn,max", "Vu≤ØVn"] if c in df_shear_all.columns]
    df_shear_all = df_shear_all.drop(columns=cols_to_drop)
    doc_builder.add_table_data(
        df_shear_all, column_widths=doc_builder.content_widths(df_shear_all), font_size=SUMMARY_FONT_SIZE
    )

    doc_builder.add_heading("Design Check Summary", level=3)
    df_check = self.check()
    doc_builder.add_table_data(
        df_check, column_widths=doc_builder.content_widths(df_check), font_size=SUMMARY_FONT_SIZE
    )

    doc_builder.save(f"Shear_Wall_Summary_{self.concrete.design_code}.docx")
    print(f"✅ Results exported to Shear_Wall_Summary_{self.concrete.design_code}.docx")
